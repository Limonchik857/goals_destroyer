"""Разбор прикреплённых к задаче файлов.

Из текста вложения извлекается полезная информация:
- даты («15.08.2026», «до 15 августа») — кандидаты в дедлайн задачи;
- пункты списков («- [ ] позвонить», «1. собрать отчёт») — кандидаты
  в новые задачи проекта.

Результат сохраняется в TaskFile.analysis один раз при загрузке, чтобы
не разбирать файл заново при каждом открытии страницы задачи.
"""
import datetime
import io
import re
from pathlib import Path

from django.utils import timezone

# Пределы разбора: защита от гигантских файлов (лимита размера загрузки нет).
MAX_TEXT_LENGTH = 200_000  # символов текста
MAX_PDF_PAGES = 50
MAX_DATES = 8
MAX_ITEMS = 30

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".log"}

MONTHS_RU = {
    "января": 1, "февраля": 2, "марта": 3, "апреля": 4,
    "мая": 5, "июня": 6, "июля": 7, "августа": 8,
    "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12,
}

# 15.08.2026, 15/08, 15.08.26
NUMERIC_DATE_RE = re.compile(r"\b(\d{1,2})[./](\d{1,2})(?:[./](\d{2,4}))?\b")
# 15 августа 2026, 15 августа
WORD_DATE_RE = re.compile(
    r"\b(\d{1,2})\s+(" + "|".join(MONTHS_RU) + r")(?:\s+(\d{4}))?\b",
    re.IGNORECASE,
)

# - [ ] пункт  /  * [ ] пункт  (отмеченные [x] считаются сделанными)
CHECKBOX_RE = re.compile(r"^\s*[-*]\s*\[\s?\]\s+(.+)$")
CHECKBOX_DONE_RE = re.compile(r"^\s*[-*]\s*\[[xXхХ]\]\s+(.+)$")
# - пункт  /  * пункт  /  • пункт  /  1. пункт  /  1) пункт
BULLET_RE = re.compile(r"^\s*(?:[-*•]|\d{1,3}[.)])\s+(.+)$")


def extract_text(task_file):
    """Достать текст из вложения. None — тип не поддерживается.

    Поддерживаются простой текст (.txt/.md/.csv/.log), .docx и .pdf.
    """
    extension = Path(task_file.original_name).suffix.lower()
    if extension in TEXT_EXTENSIONS:
        with task_file.file.open("rb") as handle:
            raw = handle.read(MAX_TEXT_LENGTH * 4)
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("cp1251", errors="replace")

    if extension == ".docx":
        import docx

        with task_file.file.open("rb") as handle:
            document = docx.Document(io.BytesIO(handle.read()))
        parts = [p.text for p in document.paragraphs]
        # Текст из таблиц: в офисных docx списки часто оформлены таблицей.
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)[:MAX_TEXT_LENGTH]

    if extension == ".pdf":
        from pypdf import PdfReader

        with task_file.file.open("rb") as handle:
            reader = PdfReader(io.BytesIO(handle.read()))
        parts = []
        total = 0
        for page in reader.pages[:MAX_PDF_PAGES]:
            text = page.extract_text() or ""
            parts.append(text)
            total += len(text)
            if total >= MAX_TEXT_LENGTH:
                break
        return "\n".join(parts)[:MAX_TEXT_LENGTH]

    return None


def _resolve_date(day, month, year, today):
    """Собрать дату из частей; None, если дата не существует или в прошлом.

    Дата без года трактуется как ближайшая будущая: «до 15 августа» в
    сентябре означает следующий год.
    """
    if year is not None and year < 100:
        year += 2000
    try:
        if year is not None:
            date = datetime.date(year, month, day)
            return date if date >= today else None
        date = datetime.date(today.year, month, day)
        if date < today:
            date = datetime.date(today.year + 1, month, day)
        return date
    except ValueError:
        return None


def find_dates(text):
    """Найти в тексте будущие даты — кандидаты в дедлайн. ISO-строки."""
    today = timezone.localdate()
    found = set()
    for match in NUMERIC_DATE_RE.finditer(text):
        day, month, year = match.groups()
        date = _resolve_date(
            int(day), int(month), int(year) if year else None, today
        )
        if date:
            found.add(date)
    for match in WORD_DATE_RE.finditer(text):
        day, month_name, year = match.groups()
        date = _resolve_date(
            int(day),
            MONTHS_RU[month_name.lower()],
            int(year) if year else None,
            today,
        )
        if date:
            found.add(date)
    return [d.isoformat() for d in sorted(found)[:MAX_DATES]]


def find_items(text):
    """Найти пункты списков — кандидаты в задачи.

    Если в тексте есть чекбоксы, берутся только неотмеченные «- [ ]»:
    это явные todo. Иначе — обычные маркированные/нумерованные пункты.
    """
    lines = text.splitlines()
    checkboxes = []
    has_checkboxes = False
    bullets = []
    for line in lines:
        if CHECKBOX_RE.match(line):
            has_checkboxes = True
            checkboxes.append(CHECKBOX_RE.match(line).group(1))
            continue
        if CHECKBOX_DONE_RE.match(line):
            has_checkboxes = True
            continue
        match = BULLET_RE.match(line)
        if match:
            bullets.append(match.group(1))

    items = checkboxes if has_checkboxes else bullets
    cleaned = []
    for item in items:
        item = item.strip()[:200]  # длина Task.name
        if len(item) >= 3:
            cleaned.append(item)
    return cleaned[:MAX_ITEMS]


def analyze_attachment(task_file):
    """Разобрать вложение и сохранить результат в task_file.analysis.

    Никогда не роняет сохранение задачи: битый или нечитаемый файл — это
    просто вложение без разбора (analysis = None).
    """
    try:
        text = extract_text(task_file)
    except Exception:
        text = None
    if text is None:
        task_file.analysis = None
    else:
        task_file.analysis = {
            "dates": find_dates(text),
            "items": find_items(text),
        }
    task_file.save(update_fields=["analysis"])
    return task_file.analysis
