import datetime
import os
import shutil
import tempfile

from django.conf import settings
from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone

from tasks.models import (
    HistoryEntry,
    JournalEntry,
    Note,
    Project,
    ProjectTemplate,
    Task,
    TaskFile,
    TemplateTask,
    add_months,
)
from tasks.services.attachment_analysis import MONTHS_RU, find_dates, find_items
from tasks.services.gamification import (
    achievements,
    level_info,
    streak_info,
    summary,
    task_points,
)


def _task_payload(**extra):
    payload = {
        "name": "Тест",
        "description": "Описание задачи",
        "project": "",
        "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
        "priority": Task.Priority.LOW,
        "difficulty": Task.Difficulty.MEDIUM,
        "estimated_duration": Task.EstimatedDuration.UP_TO_30,
        "recurrence": Task.Recurrence.NONE,
        "recurrence_interval_days": "",
    }
    payload.update(extra)
    return payload


class SmokeTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.project = Project.objects.create(owner=self.user, name="Proj")
        self.task_active = Task.objects.create(
            owner=self.user,
            name="Active",
            project=self.project,
            deadline=timezone.localdate() + timezone.timedelta(days=2),
            priority=Task.Priority.HIGH,
        )
        self.task_done = Task.objects.create(
            owner=self.user,
            name="Done",
            project=self.project,
            status=Task.Status.DONE,
        )
        self.task_no_deadline = Task.objects.create(
            owner=self.user,
            name="NoDeadline",
        )
        self.task_overdue = Task.objects.create(
            owner=self.user,
            name="Overdue",
            deadline=timezone.localdate() - timezone.timedelta(days=1),
        )

    def test_main_pages_load(self):
        urls = [
            reverse("tasks:home"),
            reverse("tasks:workspace"),
            reverse("tasks:overdue_tasks"),
            reverse("tasks:completed_tasks"),
            reverse("tasks:completed_projects"),
            reverse("tasks:history"),
            reverse("tasks:project_list"),
            reverse("tasks:project_detail", kwargs={"pk": self.project.pk}),
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            reverse("tasks:task_create"),
            reverse("tasks:task_edit", kwargs={"pk": self.task_active.pk}),
            reverse("tasks:stats"),
        ]
        for url in urls:
            with self.subTest(url=url):
                r = self.client.get(url)
                self.assertEqual(r.status_code, 200, f"GET {url}")

    def test_dashboard_overdue_card_links_to_overdue_tab(self):
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, reverse("tasks:overdue_tasks"))

    def test_dashboard_shows_next_task(self):
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, "Следующая задача")
        # У просроченной задачи самый близкий дедлайн — она и есть следующая.
        self.assertEqual(r.context["next_task"], self.task_overdue)

    def test_next_task_prefers_deadline_over_priority(self):
        # Задача без дедлайна с высоким приоритетом не обгоняет задачу
        # с ближайшим дедлайном.
        r = self.client.get(reverse("tasks:home"))
        self.assertEqual(r.context["next_task"], self.task_overdue)

    def test_next_task_fallback_to_priority_without_deadlines(self):
        Task.objects.filter(owner=self.user, deadline__isnull=False).delete()
        Task.objects.create(owner=self.user, name="LowPrio",
                            priority=Task.Priority.LOW)
        high = Task.objects.create(owner=self.user, name="HighPrio",
                                   priority=Task.Priority.HIGH)
        r = self.client.get(reverse("tasks:home"))
        self.assertEqual(r.context["next_task"], high)

    def test_next_task_fallback_to_earliest_when_all_equal(self):
        # Нет ни дедлайнов, ни различий в приоритетах → раньше добавленная.
        Task.objects.filter(owner=self.user).delete()
        first = Task.objects.create(owner=self.user, name="First")
        Task.objects.create(owner=self.user, name="Second")
        r = self.client.get(reverse("tasks:home"))
        self.assertEqual(r.context["next_task"], first)

    def test_next_task_empty_state_when_no_active_tasks(self):
        Task.objects.filter(owner=self.user).delete()
        r = self.client.get(reverse("tasks:home"))
        self.assertIsNone(r.context["next_task"])
        # Раздел не пропадает: заголовок на месте, вместо карточки заглушка.
        self.assertContains(r, "Следующая задача")
        self.assertContains(r, "Задач нет")

    def test_workspace_shows_priority_and_deadline_label(self):
        r = self.client.get(reverse("tasks:workspace"))
        self.assertContains(r, "Высокий")
        self.assertContains(r, "Без дедлайна")
        self.assertContains(r, "Проект: Proj")

    def test_workspace_overdue_first(self):
        r = self.client.get(reverse("tasks:workspace"))
        content = r.content.decode()
        # Просроченная задача идёт раньше остальных активных в выдаче.
        self.assertLess(
            content.index("Overdue"), content.index("NoDeadline")
        )
        # Вкладки-переключатели на месте.
        self.assertContains(r, "subtabs")

    def test_overdue_page_shows_only_overdue(self):
        r = self.client.get(reverse("tasks:overdue_tasks"))
        self.assertContains(r, self.task_overdue.name)
        self.assertNotContains(r, self.task_active.name)
        self.assertNotContains(r, self.task_no_deadline.name)

    def test_project_detail_title_and_back_link(self):
        r = self.client.get(
            reverse("tasks:project_detail", kwargs={"pk": self.project.pk})
        )
        self.assertContains(r, "panel__title\">Proj")
        self.assertContains(r, "← Все проекты")
        # На странице проекта видны и активные, и завершённые задачи.
        self.assertContains(r, self.task_active.name)
        self.assertContains(r, self.task_done.name)

    def test_project_list_has_complete_checkbox(self):
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(r, 'aria-label="Завершить проект"')
        self.assertNotContains(r, ">Завершить</a>")

    def test_completed_tasks_page(self):
        r = self.client.get(reverse("tasks:completed_tasks"))
        self.assertContains(r, self.task_done.name)
        self.assertNotContains(r, self.task_active.name)

    def test_completed_projects_page(self):
        self.project.status = Project.Status.COMPLETED
        self.project.save()
        r = self.client.get(reverse("tasks:completed_projects"))
        self.assertContains(r, self.project.name)
        # А из списка активных проектов он пропал.
        r2 = self.client.get(reverse("tasks:project_list"))
        self.assertNotContains(r2, self.project.name)

    def test_task_toggle_moves_to_completed(self):
        r = self.client.post(
            reverse("tasks:task_toggle", kwargs={"pk": self.task_active.pk}),
            {"next": reverse("tasks:workspace")},
        )
        self.assertEqual(r.status_code, 302)
        self.assertTrue(Task.objects.get(pk=self.task_active.pk).is_done)

    def test_project_complete_with_pending_tasks(self):
        r = self.client.post(
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            {"action": "complete"},
        )
        self.assertEqual(r.status_code, 302)
        proj = Project.objects.get(pk=self.project.pk)
        self.assertTrue(proj.is_completed)
        self.assertTrue(Task.objects.get(pk=self.task_active.pk).is_done)

    def test_project_complete_cancel_does_nothing(self):
        r = self.client.post(
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            {"action": "cancel"},
        )
        self.assertEqual(r.status_code, 302)
        self.assertFalse(Project.objects.get(pk=self.project.pk).is_completed)
        self.assertFalse(Task.objects.get(pk=self.task_active.pk).is_done)

    def test_project_complete_confirm_lists_pending(self):
        r = self.client.get(
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk})
        )
        self.assertContains(r, "невыполненн")
        self.assertContains(r, self.task_active.name)

    def test_project_reopen(self):
        self.project.status = Project.Status.COMPLETED
        self.project.save()
        r = self.client.post(
            reverse("tasks:project_reopen", kwargs={"pk": self.project.pk})
        )
        self.assertEqual(r.status_code, 302)
        self.assertFalse(Project.objects.get(pk=self.project.pk).is_completed)


class CalendarTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.today = timezone.localdate()
        Task.objects.create(owner=self.user, name="T1", deadline=self.today)
        Task.objects.create(owner=self.user, name="T2", deadline=self.today)
        self.other = Task.objects.create(
            owner=self.user,
            name="Other",
            deadline=self.today + timezone.timedelta(days=3),
        )

    def test_calendar_page_loads(self):
        r = self.client.get(reverse("tasks:calendar"))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "просмотреть")

    def test_calendar_shows_deadline_count(self):
        r = self.client.get(reverse("tasks:calendar"))
        # На сегодня два дедлайна — текст "2 дедлайна".
        self.assertContains(r, "calendar-day__count\">2 дедлайна")

    def test_calendar_done_tasks_not_counted(self):
        Task.objects.create(
            owner=self.user,
            name="DoneOne",
            deadline=self.today,
            status=Task.Status.DONE,
        )
        r = self.client.get(reverse("tasks:calendar"))
        self.assertContains(r, "calendar-day__count\">2 дедлайна")

    def test_calendar_month_navigation(self):
        r = self.client.get(reverse("tasks:calendar") + "?year=2020&month=1")
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Январь 2020")

    def test_calendar_invalid_params_redirect(self):
        r = self.client.get(reverse("tasks:calendar") + "?year=abc&month=99")
        self.assertEqual(r.status_code, 302)

    def test_calendar_day_page(self):
        url = reverse(
            "tasks:calendar_day",
            kwargs={
                "year": self.today.year,
                "month": self.today.month,
                "day": self.today.day,
            },
        )
        r = self.client.get(url)
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "T1")
        self.assertContains(r, "T2")
        self.assertNotContains(r, "Other")

    def test_calendar_day_invalid_date_404(self):
        url = reverse(
            "tasks:calendar_day",
            kwargs={"year": 2026, "month": 2, "day": 30},
        )
        r = self.client.get(url)
        self.assertEqual(r.status_code, 404)


class DashboardButtonTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_dashboard_shows_notes_and_stats(self):
        r = self.client.get(reverse("tasks:home"))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Активные задачи")
        self.assertContains(r, "Последние заметки")


class DashboardFocusSectionTest(TestCase):
    """Секция «По фокусу и энергии» на главной."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_offers_focus_when_no_recent_session(self):
        # Оценки ещё не было → предложение перейти в «Фокус и энергию».
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, "По фокусу и энергии")
        self.assertContains(r, "Оцените свою энергию и концентрацию")
        self.assertIsNone(r.context["focus_task"])
        self.assertFalse(r.context["focus_recent"])

    def test_old_session_ignored(self):
        # Оценка старше 24 часов не учитывается.
        from focus.models import WorkSession

        WorkSession.objects.create(user=self.user, energy=2, focus=2, available_time=2)
        WorkSession.objects.filter(user=self.user).update(
            created_at=timezone.now() - datetime.timedelta(hours=25)
        )
        r = self.client.get(reverse("tasks:home"))
        self.assertFalse(r.context["focus_recent"])
        self.assertIsNone(r.context["focus_task"])
        self.assertContains(r, "Оцените свою энергию и концентрацию")

    def test_recent_session_shows_recommended_task(self):
        from focus.models import WorkSession

        Task.objects.create(owner=self.user, name="LowPrio",
                            priority=Task.Priority.LOW)
        high = Task.objects.create(owner=self.user, name="HighPrio",
                                   priority=Task.Priority.HIGH)
        WorkSession.objects.create(
            user=self.user, energy=2, focus=2, available_time=2,
        )
        r = self.client.get(reverse("tasks:home"))
        self.assertTrue(r.context["focus_recent"])
        self.assertEqual(r.context["focus_task"], high)

    def test_new_task_appears_after_next_closed(self):
        from focus.models import WorkSession

        first = Task.objects.create(owner=self.user, name="First",
                                    priority=Task.Priority.HIGH)
        Task.objects.create(owner=self.user, name="Second",
                            priority=Task.Priority.LOW)
        WorkSession.objects.create(
            user=self.user, energy=2, focus=2, available_time=2,
        )
        r = self.client.get(reverse("tasks:home"))
        self.assertEqual(r.context["focus_task"], first)

        # Закрываем рекомендацию — на её место встаёт следующая.
        first.status = Task.Status.DONE
        first.save()
        r = self.client.get(reverse("tasks:home"))
        self.assertEqual(r.context["focus_task"].name, "Second")

    def test_no_tasks_shows_empty_state(self):
        from focus.models import WorkSession

        WorkSession.objects.create(
            user=self.user, energy=2, focus=2, available_time=2,
        )
        r = self.client.get(reverse("tasks:home"))
        self.assertTrue(r.context["focus_recent"])
        self.assertIsNone(r.context["focus_task"])
        self.assertContains(r, "Задач нет")

    def test_recommended_task_card_rendered(self):
        from focus.models import WorkSession

        Task.objects.create(owner=self.user, name="FocusCard",
                            priority=Task.Priority.HIGH)
        WorkSession.objects.create(
            user=self.user, energy=2, focus=2, available_time=2,
        )
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, "FocusCard")


class DashboardTabsTest(TestCase):
    """Переключатель в секции «Следующая задача»."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_next_task_section_has_switcher(self):
        Task.objects.create(owner=self.user, name="Задача",
                            priority=Task.Priority.HIGH)
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, "По приоритету и дедлайну")
        self.assertContains(r, "По фокусу и энергии")
        self.assertContains(r, 'id="next-task-tabs"')
        self.assertContains(r, 'id="next-priority"')
        self.assertContains(r, 'id="next-focus"')

    def test_focus_panel_offers_link_without_session(self):
        Task.objects.create(owner=self.user, name="Задача",
                            priority=Task.Priority.HIGH)
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, "Оцените свою энергию и концентрацию")


class WorkspaceRecurringButtonTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_workspace_has_recurring_button(self):
        r = self.client.get(reverse("tasks:workspace"))
        self.assertContains(r, "Повторяющаяся задача")
        self.assertContains(
            r, reverse("tasks:task_create") + "?recurring=1"
        )

    def test_recurring_param_preselects_recurrence(self):
        r = self.client.get(reverse("tasks:task_create") + "?recurring=1")
        form = r.context["form"]
        self.assertEqual(
            form["recurrence"].value(), Task.Recurrence.EVERY_N_DAYS
        )
        self.assertEqual(form["recurrence_interval_days"].value(), 7)

    def test_recurring_task_creation_works(self):
        r = self.client.post(
            reverse("tasks:task_create") + "?recurring=1",
            {"name": "Отчёт", "description": "Еженедельный отчёт", "project": "",
             "deadline": timezone.localdate().isoformat(),
             "priority": Task.Priority.LOW,
             "difficulty": Task.Difficulty.MEDIUM,
             "estimated_duration": Task.EstimatedDuration.UP_TO_30,
             "recurrence": Task.Recurrence.EVERY_N_DAYS,
             "recurrence_interval_days": "7"},
        )
        self.assertEqual(r.status_code, 302)
        task = Task.objects.get(owner=self.user, name="Отчёт")
        self.assertEqual(task.recurrence, Task.Recurrence.EVERY_N_DAYS)
        self.assertEqual(task.recurrence_interval_days, 7)

    def test_recurrence_select_still_has_every_n_days(self):
        r = self.client.get(reverse("tasks:task_create"))
        form = r.context["form"]
        self.assertIn(
            Task.Recurrence.EVERY_N_DAYS,
            [v for v, _ in form["recurrence"].field.choices],
        )

    def test_clearing_interval_disables_recurrence(self):
        # Редактирование повторяющейся задачи: убираем интервал —
        # повторение отключается.
        task = Task.objects.create(
            owner=self.user, name="Отчёт",
            recurrence=Task.Recurrence.EVERY_N_DAYS,
            recurrence_interval_days=7,
        )
        r = self.client.post(
            reverse("tasks:task_edit", kwargs={"pk": task.pk}),
            {"name": "Отчёт", "description": "Описание", "project": "",
             "deadline": timezone.localdate().isoformat(),
             "priority": Task.Priority.LOW,
             "difficulty": Task.Difficulty.MEDIUM,
             "estimated_duration": Task.EstimatedDuration.UP_TO_30,
             "recurrence": Task.Recurrence.NONE,
             "recurrence_interval_days": ""},
        )
        self.assertEqual(r.status_code, 302)
        task.refresh_from_db()
        self.assertEqual(task.recurrence, Task.Recurrence.NONE)
        self.assertIsNone(task.recurrence_interval_days)


class HistoryTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.project = Project.objects.create(owner=self.user, name="Proj")

    def test_task_actions_are_logged(self):
        self.client.post(
            reverse("tasks:task_create"),
            {"name": "NewTask", "description": "Описание", "project": "",
             "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
             "priority": Task.Priority.MEDIUM,
             "difficulty": Task.Difficulty.MEDIUM,
             "estimated_duration": Task.EstimatedDuration.UP_TO_30,
             "recurrence": Task.Recurrence.NONE},
        )
        task = Task.objects.get(owner=self.user, name="NewTask")
        self.client.post(reverse("tasks:task_toggle", kwargs={"pk": task.pk}))
        texts = list(
            HistoryEntry.objects.filter(owner=self.user).values_list(
                "text", flat=True
            )
        )
        self.assertIn("Создана задача «NewTask»", texts)
        self.assertIn("Задача «NewTask» выполнена", texts)

    def test_project_complete_and_reopen_are_logged(self):
        self.client.post(
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            {"action": "complete"},
        )
        self.client.post(
            reverse("tasks:project_reopen", kwargs={"pk": self.project.pk})
        )
        texts = list(
            HistoryEntry.objects.filter(owner=self.user).values_list(
                "text", flat=True
            )
        )
        self.assertIn("Проект «Proj» завершён", texts)
        self.assertIn("Проект «Proj» возвращён в активные", texts)

    def test_history_page_shows_day_and_time(self):
        self.client.post(
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            {"action": "complete"},
        )
        r = self.client.get(reverse("tasks:history"))
        self.assertContains(r, "Проект «Proj» завершён")
        # День жирным заголовком и время действия.
        self.assertContains(r, "history-day__title")
        self.assertContains(r, timezone.localtime().strftime("%H:%M"))


class ProjectTemplateTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.template = ProjectTemplate.objects.create(
            owner=self.user, name="Сайт", description="Типовой сайт"
        )
        self.tt1 = TemplateTask.objects.create(
            template=self.template,
            name="Купить домен",
            description="На reg.ru",
            priority=Task.Priority.HIGH,
            deadline_offset_days=7,
        )
        self.tt2 = TemplateTask.objects.create(
            template=self.template,
            name="Настроить Django",
            priority=Task.Priority.LOW,
        )

    def test_template_pages_load(self):
        urls = [
            reverse("tasks:template_list"),
            reverse("tasks:template_create"),
            reverse("tasks:template_detail", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_edit", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_delete", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_task_create", kwargs={"pk": self.template.pk}),
            reverse(
                "tasks:template_task_edit",
                kwargs={"pk": self.template.pk, "task_pk": self.tt1.pk},
            ),
            reverse(
                "tasks:template_task_delete",
                kwargs={"pk": self.template.pk, "task_pk": self.tt1.pk},
            ),
        ]
        for url in urls:
            with self.subTest(url=url):
                r = self.client.get(url)
                self.assertEqual(r.status_code, 200, f"GET {url}")

    def test_template_detail_lists_tasks(self):
        r = self.client.get(
            reverse("tasks:template_detail", kwargs={"pk": self.template.pk})
        )
        self.assertContains(r, "Купить домен")
        self.assertContains(r, "Настроить Django")
        self.assertContains(r, "Высокий")

    def test_add_task_to_template(self):
        r = self.client.post(
            reverse("tasks:template_task_create", kwargs={"pk": self.template.pk}),
            {"name": "Сделать деплой", "description": "",
             "priority": Task.Priority.MEDIUM, "deadline_offset_days": ""},
        )
        self.assertEqual(r.status_code, 302)
        self.assertTrue(
            TemplateTask.objects.filter(
                template=self.template, name="Сделать деплой"
            ).exists()
        )

    def test_create_project_from_template_copies_everything(self):
        today = timezone.localdate()
        r = self.client.post(
            reverse("tasks:project_create"),
            {"template": str(self.template.pk), "name": "",
             "description": "Типовой сайт",
             "deadline": (today + timezone.timedelta(days=14)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        self.assertEqual(r.status_code, 302)
        project = Project.objects.get(owner=self.user, name="Сайт")
        # Название и описание взялись из шаблона.
        self.assertEqual(project.description, "Типовой сайт")
        tasks = {t.name: t for t in project.tasks.all()}
        self.assertEqual(set(tasks), {"Купить домен", "Настроить Django"})
        # Приоритеты, описания и дедлайн-смещение скопированы.
        self.assertEqual(tasks["Купить домен"].priority, Task.Priority.HIGH)
        self.assertEqual(tasks["Купить домен"].description, "На reg.ru")
        self.assertEqual(
            tasks["Купить домен"].deadline,
            today + timezone.timedelta(days=7),
        )
        self.assertEqual(tasks["Настроить Django"].priority, Task.Priority.LOW)
        self.assertIsNone(tasks["Настроить Django"].deadline)

    def test_created_project_is_independent_from_template(self):
        today = timezone.localdate()
        self.client.post(
            reverse("tasks:project_create"),
            {"template": str(self.template.pk), "name": "",
             "description": "Типовой сайт",
             "deadline": (today + timezone.timedelta(days=14)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        project = Project.objects.get(owner=self.user, name="Сайт")
        # Изменение шаблона не влияет на уже созданный проект.
        self.tt1.name = "Переименованная"
        self.tt1.save()
        self.tt2.delete()
        names = set(project.tasks.values_list("name", flat=True))
        self.assertEqual(names, {"Купить домен", "Настроить Django"})

    def test_create_project_from_template_with_custom_name(self):
        today = timezone.localdate()
        self.client.post(
            reverse("tasks:project_create"),
            {"template": str(self.template.pk), "name": "Магазин",
             "description": "Типовой сайт",
             "deadline": (today + timezone.timedelta(days=14)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        self.assertTrue(
            Project.objects.filter(owner=self.user, name="Магазин").exists()
        )

    def test_create_empty_project_still_works(self):
        r = self.client.post(
            reverse("tasks:project_create"),
            {"template": "", "name": "Пустой", "description": "Описание",
             "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        self.assertEqual(r.status_code, 302)
        project = Project.objects.get(owner=self.user, name="Пустой")
        self.assertEqual(project.tasks.count(), 0)

    def test_project_create_requires_name_without_template(self):
        r = self.client.post(
            reverse("tasks:project_create"),
            {"template": "", "name": "", "description": "Описание",
             "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        self.assertEqual(r.status_code, 200)  # форма с ошибкой, без редиректа
        self.assertEqual(Project.objects.count(), 0)

    def test_template_creation_logged(self):
        today = timezone.localdate()
        self.client.post(
            reverse("tasks:project_create"),
            {"template": str(self.template.pk), "name": "",
             "description": "Типовой сайт",
             "deadline": (today + timezone.timedelta(days=14)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        self.assertTrue(
            HistoryEntry.objects.filter(
                owner=self.user, text="Создан проект «Сайт» по шаблону «Сайт»"
            ).exists()
        )


class ProjectTemplatesSectionTest(TestCase):
    """Секция шаблонов на странице проектов и диалог создания по шаблону."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_project_page_lists_templates(self):
        ProjectTemplate.objects.create(owner=self.user, name="Онбординг")
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(r, "Онбординг")

    def test_project_page_has_create_from_template_button(self):
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(r, "Создать по шаблону")
        self.assertContains(r, 'id="open-template-dialog"')

    def test_template_dialog_empty_state(self):
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(r, "Шаблонов нет")

    def test_template_dialog_lists_templates(self):
        t = ProjectTemplate.objects.create(owner=self.user, name="Онбординг")
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(
            r, reverse("tasks:project_create") + "?template=" + str(t.pk)
        )

    def test_new_template_button_visible(self):
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(r, "Создать шаблон")
        self.assertContains(r, reverse("tasks:template_create"))

    def test_templates_not_in_top_nav(self):
        r = self.client.get(reverse("tasks:project_list"))
        # Верхнее меню (navbar-sub) не содержит ссылки на список шаблонов.
        nav = r.content.decode()
        self.assertNotIn(">Шаблоны</a>", nav)

    def test_dialog_has_edit_links_for_templates(self):
        t = ProjectTemplate.objects.create(owner=self.user, name="Онбординг")
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(
            r, reverse("tasks:template_edit", kwargs={"pk": t.pk})
        )


class ProjectCreateFromDialogTest(TestCase):
    """Создание проекта после выбора шаблона в диалоге."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.template = ProjectTemplate.objects.create(
            owner=self.user, name="Сайт", description="Типовой сайт"
        )
        self.step = TemplateTask.objects.create(
            template=self.template, name="Купить домен",
            deadline_offset_days=7,
        )

    def test_template_field_hidden_when_preselected(self):
        r = self.client.get(
            reverse("tasks:project_create") + "?template=" + str(self.template.pk)
        )
        self.assertContains(r, "создаётся из шаблона «Сайт»")
        # Вместо выпадающего списка — скрытое поле с выбранным шаблоном.
        self.assertContains(r, 'type="hidden" name="template"')
        self.assertNotContains(r, "— Пустой проект —")

    def test_preselected_template_survives_post(self):
        r = self.client.post(
            reverse("tasks:project_create") + "?template=" + str(self.template.pk),
            {"template": str(self.template.pk), "name": "", "description": "Типовой сайт",
             "deadline": (timezone.localdate() + timezone.timedelta(days=14)).isoformat(),
             "tasks-TOTAL_FORMS": "0", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000"},
        )
        self.assertEqual(r.status_code, 302)
        self.assertTrue(Project.objects.filter(owner=self.user, name="Сайт").exists())

    def test_form_prefilled_from_template(self):
        today = timezone.localdate()
        r = self.client.get(
            reverse("tasks:project_create") + "?template=" + str(self.template.pk)
        )
        form = r.context["form"]
        self.assertEqual(form["name"].value(), "Сайт")
        self.assertEqual(form["description"].value(), "Типовой сайт")
        self.assertEqual(
            form["deadline"].value(), today + timezone.timedelta(days=7)
        )

    def test_project_create_page_still_has_template_selector(self):
        # Без выбранного шаблона поле выбора остаётся на месте.
        r = self.client.get(reverse("tasks:project_create"))
        self.assertContains(r, 'name="template"')

    def test_inline_task_block_has_difficulty_and_duration(self):
        r = self.client.get(reverse("tasks:project_create"))
        self.assertContains(r, "Сложность")
        self.assertContains(r, "Длительность")
        self.assertContains(r, 'name="tasks-0-difficulty"')
        self.assertContains(r, 'name="tasks-0-estimated_duration"')

    def test_project_with_inline_tasks_created(self):
        r = self.client.post(
            reverse("tasks:project_create"),
            {"template": "", "name": "Проект", "description": "Описание",
             "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
             "tasks-TOTAL_FORMS": "1", "tasks-INITIAL_FORMS": "0",
             "tasks-MIN_NUM_FORMS": "0", "tasks-MAX_NUM_FORMS": "1000",
             "tasks-0-name": "Задача 1",
             "tasks-0-priority": Task.Priority.MEDIUM,
             "tasks-0-difficulty": Task.Difficulty.HARD,
             "tasks-0-estimated_duration": Task.EstimatedDuration.OVER_60,
             "tasks-0-deadline": (timezone.localdate() + timezone.timedelta(days=3)).isoformat(),
             "tasks-0-description": "Сделать",
             "tasks-0-DELETE": ""},
        )
        self.assertEqual(r.status_code, 302)
        task = Task.objects.get(owner=self.user, name="Задача 1")
        self.assertEqual(task.difficulty, Task.Difficulty.HARD)
        self.assertEqual(task.estimated_duration, Task.EstimatedDuration.OVER_60)


class TemplateFormWithStepsTest(TestCase):
    """Создание/редактирование шаблона с этапами прямо в форме."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def _payload(self, name="Онбординг", steps=()):
        data = {
            "name": name,
            "description": "Типовой онбординг",
            "tasks-TOTAL_FORMS": str(len(steps)),
            "tasks-INITIAL_FORMS": "0",
            "tasks-MIN_NUM_FORMS": "0",
            "tasks-MAX_NUM_FORMS": "1000",
        }
        for i, step in enumerate(steps):
            prefix = f"tasks-{i}-"
            data.update({
                f"{prefix}name": step.get("name", ""),
                f"{prefix}priority": step.get("priority", Task.Priority.MEDIUM),
                f"{prefix}difficulty": step.get(
                    "difficulty", Task.Difficulty.MEDIUM
                ),
                f"{prefix}estimated_duration": step.get(
                    "estimated_duration", Task.EstimatedDuration.UP_TO_30
                ),
                f"{prefix}deadline_offset_days": step.get(
                    "deadline_offset_days", ""
                ),
                f"{prefix}description": step.get("description", ""),
            })
        return data

    def test_create_template_with_steps(self):
        r = self.client.post(
            reverse("tasks:template_create"),
            self._payload(steps=[
                {"name": "Завести аккаунт", "priority": Task.Priority.HIGH,
                 "difficulty": Task.Difficulty.EASY,
                 "estimated_duration": Task.EstimatedDuration.UP_TO_15,
                 "deadline_offset_days": "1"},
                {"name": "Настроить окружение"},
            ]),
        )
        self.assertEqual(r.status_code, 302)
        template = ProjectTemplate.objects.get(owner=self.user, name="Онбординг")
        steps = {s.name: s for s in template.template_tasks.all()}
        self.assertEqual(set(steps), {"Завести аккаунт", "Настроить окружение"})
        self.assertEqual(steps["Завести аккаунт"].priority, Task.Priority.HIGH)
        self.assertEqual(steps["Завести аккаунт"].difficulty, Task.Difficulty.EASY)
        self.assertEqual(steps["Завести аккаунт"].deadline_offset_days, 1)

    def test_create_template_without_steps(self):
        r = self.client.post(
            reverse("tasks:template_create"),
            self._payload(steps=[]),
        )
        self.assertEqual(r.status_code, 302)
        template = ProjectTemplate.objects.get(owner=self.user, name="Онбординг")
        self.assertEqual(template.template_tasks.count(), 0)

    def test_create_form_has_step_fields(self):
        r = self.client.get(reverse("tasks:template_create"))
        self.assertContains(r, "Этапы шаблона")
        self.assertContains(r, "Сложность")
        self.assertContains(r, "Длительность")
        self.assertContains(r, "Создать шаблон проекта")

    def test_update_template_replaces_steps(self):
        template = ProjectTemplate.objects.create(
            owner=self.user, name="Онбординг"
        )
        TemplateTask.objects.create(
            template=template, name="Старый шаг"
        )
        r = self.client.post(
            reverse("tasks:template_edit", kwargs={"pk": template.pk}),
            {
                "name": "Онбординг",
                "description": "Типовой онбординг",
                "tasks-TOTAL_FORMS": "1",
                "tasks-INITIAL_FORMS": "1",
                "tasks-MIN_NUM_FORMS": "0",
                "tasks-MAX_NUM_FORMS": "1000",
                "tasks-0-id": str(
                    template.template_tasks.first().pk
                ),
                "tasks-0-name": "Новый шаг",
                "tasks-0-priority": Task.Priority.LOW,
                "tasks-0-difficulty": Task.Difficulty.MEDIUM,
                "tasks-0-estimated_duration": Task.EstimatedDuration.UP_TO_30,
                "tasks-0-deadline_offset_days": "",
                "tasks-0-description": "",
                "tasks-0-DELETE": "",
            },
        )
        self.assertEqual(r.status_code, 302)
        names = list(
            template.template_tasks.values_list("name", flat=True)
        )
        self.assertEqual(names, ["Новый шаг"])


class RecurrenceTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.today = timezone.localdate()

    def _make(self, recurrence, **kwargs):
        return Task.objects.create(
            owner=self.user,
            name="Отчёт",
            recurrence=recurrence,
            **kwargs,
        )

    def _toggle(self, task):
        return self.client.post(
            reverse("tasks:task_toggle", kwargs={"pk": task.pk})
        )

    def test_daily_spawns_next_day(self):
        task = self._make(Task.Recurrence.DAILY, deadline=self.today)
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertEqual(new.deadline, self.today + timezone.timedelta(days=1))

    def test_weekly_spawns_next_week(self):
        task = self._make(Task.Recurrence.WEEKLY, deadline=self.today)
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertEqual(new.deadline, self.today + timezone.timedelta(weeks=1))

    def test_monthly_clamps_month_end(self):
        import datetime as dt

        # 31 января + месяц → 28 февраля, а не 3 марта.
        self.assertEqual(add_months(dt.date(2026, 1, 31), 1), dt.date(2026, 2, 28))
        # Сдвиг от исходной даты: 31 января + 2 месяца → снова 31 марта,
        # клэмп февраля не «съедает» день месяца.
        self.assertEqual(add_months(dt.date(2026, 1, 31), 2), dt.date(2026, 3, 31))

    def test_monthly_spawns_next_month(self):
        task = self._make(Task.Recurrence.MONTHLY, deadline=self.today)
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertEqual(new.deadline, add_months(self.today, 1))

    def test_overdue_daily_spawns_tomorrow(self):
        # Ежедневная задача, закрытая с опозданием на 10 дней, не должна
        # порождать копию, просроченную на 9 дней.
        task = self._make(
            Task.Recurrence.DAILY,
            deadline=self.today - timezone.timedelta(days=10),
        )
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertEqual(new.deadline, self.today + timezone.timedelta(days=1))

    def test_overdue_weekly_keeps_weekday(self):
        deadline = self.today - timezone.timedelta(days=17)
        task = self._make(Task.Recurrence.WEEKLY, deadline=deadline)
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertGreater(new.deadline, self.today)  # в будущем
        self.assertLessEqual(
            new.deadline, self.today + timezone.timedelta(days=7)
        )
        # Ритм сохранён: тот же день недели, что у исходного дедлайна.
        self.assertEqual(new.deadline.weekday(), deadline.weekday())

    def test_double_toggle_does_not_duplicate_spawn(self):
        task = self._make(Task.Recurrence.DAILY, deadline=self.today)
        self._toggle(task)   # выполнена → появилась копия
        self._toggle(task)   # передумал, вернул в активные
        self._toggle(task)   # снова выполнена — вторая копия не нужна
        self.assertEqual(
            Task.objects.filter(
                owner=self.user, status=Task.Status.NOT_DONE
            ).count(),
            1,
        )

    def test_yearly_spawns_next_year(self):
        task = self._make(Task.Recurrence.YEARLY, deadline=self.today)
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertEqual(new.deadline.year, self.today.year + 1)
        self.assertEqual(
            (new.deadline.month, new.deadline.day),
            (self.today.month, self.today.day),
        )

    def test_none_recurrence_spawns_nothing(self):
        task = self._make(Task.Recurrence.NONE, deadline=self.today)
        self._toggle(task)
        self.assertEqual(
            Task.objects.filter(owner=self.user).count(), 1
        )

    def test_old_task_stays_done_new_keeps_fields(self):
        project = Project.objects.create(owner=self.user, name="P")
        task = self._make(
            Task.Recurrence.WEEKLY,
            deadline=self.today,
            description="Каждый понедельник",
            priority=Task.Priority.HIGH,
            project=project,
        )
        self._toggle(task)
        task.refresh_from_db()
        # Выполненная задача не изменилась обратно — история сохраняется.
        self.assertTrue(task.is_done)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertNotEqual(new.pk, task.pk)
        self.assertEqual(new.name, task.name)
        self.assertEqual(new.description, task.description)
        self.assertEqual(new.project, task.project)
        self.assertEqual(new.priority, task.priority)
        self.assertEqual(new.recurrence, task.recurrence)

    def test_recurrence_without_deadline_spawns_without_deadline(self):
        task = self._make(Task.Recurrence.DAILY)
        self._toggle(task)
        new = Task.objects.get(owner=self.user, status=Task.Status.NOT_DONE)
        self.assertIsNone(new.deadline)

    def test_uncomplete_does_not_spawn(self):
        task = self._make(Task.Recurrence.DAILY, deadline=self.today,
                          status=Task.Status.DONE)
        self._toggle(task)  # возврат в активные
        self.assertEqual(Task.objects.filter(owner=self.user).count(), 1)

    def test_recurrence_creation_logged(self):
        task = self._make(Task.Recurrence.DAILY, deadline=self.today)
        self._toggle(task)
        texts = list(
            HistoryEntry.objects.filter(owner=self.user).values_list(
                "text", flat=True
            )
        )
        self.assertIn("Задача «Отчёт» выполнена", texts)
        self.assertIn("Задача «Отчёт» создана повторно", texts)

    def test_task_form_has_recurrence_field(self):
        r = self.client.get(reverse("tasks:task_create"))
        self.assertContains(r, "Повторение")
        self.assertContains(r, "Каждую неделю")


class ImportTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.project = Project.objects.create(owner=self.user, name="Proj")

    def test_import_page_requires_login(self):
        self.client.logout()
        r = self.client.get(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk})
        )
        # Неавторизованный перенаправляется на логин.
        self.assertNotEqual(r.status_code, 200)

    def test_import_page_loads(self):
        r = self.client.get(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk})
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Импорт задач из файла")
        self.assertContains(r, "Проект: Proj")

    def test_import_other_users_project_404(self):
        other = User.objects.create_user("other", password="p")
        project = Project.objects.create(owner=other, name="Secret")
        r = self.client.get(
            reverse("tasks:import_tasks", kwargs={"pk": project.pk})
        )
        self.assertEqual(r.status_code, 404)

    def test_import_creates_tasks_from_plain_lines(self):
        content = "Задача 1\n\nЗадача 2\nЗадача 3\n".encode("utf-8")
        file = SimpleUploadedFile("tasks.txt", content, content_type="text/plain")
        r = self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"file": file,
             "default_priority": Task.Priority.LOW,
             "default_deadline": ""},
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Найдено")
        self.assertContains(r, "3 задачи")
        self.assertContains(r, "Задача 1")
        self.assertContains(r, "Задача 3")

    def test_import_parses_markdown_headings(self):
        content = "# Купить домен\nНа reg.ru\n\n# Настроить сервер\nUbuntu 24.04".encode("utf-8")
        file = SimpleUploadedFile("tasks.md", content, content_type="text/markdown")
        r = self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"file": file,
             "default_priority": Task.Priority.LOW,
             "default_deadline": ""},
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Купить домен")
        self.assertContains(r, "На reg.ru")
        self.assertContains(r, "Настроить сервер")

    def test_import_parses_checkboxes(self):
        content = "- [ ] Сделать отчёт\n- [ ] Отправить клиенту".encode("utf-8")
        file = SimpleUploadedFile("tasks.md", content, content_type="text/markdown")
        r = self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"file": file,
             "default_priority": Task.Priority.LOW,
             "default_deadline": ""},
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Сделать отчёт")
        self.assertContains(r, "Отправить клиенту")

    def test_import_confirm_creates_tasks(self):
        content = "Задача A\nЗадача B".encode("utf-8")
        file = SimpleUploadedFile("tasks.txt", content, content_type="text/plain")
        # Шаг 1: разбор
        self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"file": file,
             "default_priority": Task.Priority.HIGH,
             "default_deadline": ""},
        )
        # Шаг 2: подтверждение
        r = self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"confirm": "1"},
        )
        self.assertRedirects(r, reverse("tasks:project_detail", kwargs={"pk": self.project.pk}))
        tasks = Task.objects.filter(project=self.project).order_by("pk")
        self.assertEqual(tasks.count(), 2)
        self.assertEqual(tasks[0].name, "Задача A")
        self.assertEqual(tasks[1].name, "Задача B")
        self.assertEqual(tasks[0].priority, Task.Priority.HIGH)

    def test_import_empty_file_shows_error(self):
        file = SimpleUploadedFile("empty.txt", b"", content_type="text/plain")
        r = self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"file": file,
             "default_priority": Task.Priority.LOW,
             "default_deadline": ""},
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Отправленный файл пуст")

    def test_import_logs_history(self):
        content = "Тестовая задача".encode("utf-8")
        file = SimpleUploadedFile("tasks.txt", content, content_type="text/plain")
        self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"file": file,
             "default_priority": Task.Priority.LOW,
             "default_deadline": ""},
        )
        self.client.post(
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            {"confirm": "1"},
        )
        self.assertTrue(
            HistoryEntry.objects.filter(
                owner=self.user,
                text__contains="Импортировано 1 задач в проект «Proj»",
            ).exists()
        )


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="taskfiles-test-"))
class TaskFileTest(TestCase):
    """Прикрепление файлов к задачам: загрузка, доступ, открепление."""

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def task_payload(self, **extra):
        """Минимальный валидный POST для формы задачи."""
        payload = {
            "name": "Задача с файлами",
            "description": "Описание",
            "project": "",
            "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
            "priority": Task.Priority.LOW,
            "difficulty": Task.Difficulty.MEDIUM,
            "estimated_duration": Task.EstimatedDuration.UP_TO_30,
            "recurrence": Task.Recurrence.NONE,
            "recurrence_interval_days": "",
        }
        payload.update(extra)
        return payload

    def upload(self, name, content=b"data"):
        return SimpleUploadedFile(name, content, content_type="text/plain")

    def make_task_with_file(self, name="doc.txt"):
        task = Task.objects.create(owner=self.user, name="Задача")
        task_file = TaskFile.objects.create(
            task=task, file=self.upload(name), original_name=name
        )
        return task, task_file

    # --- Загрузка ---

    def test_create_task_attaches_several_files_at_once(self):
        self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(
                attachments=[self.upload("a.txt"), self.upload("b.txt")]
            ),
        )
        task = Task.objects.get(name="Задача с файлами")
        self.assertEqual(
            sorted(task.files.values_list("original_name", flat=True)),
            ["a.txt", "b.txt"],
        )

    def test_create_task_without_files_works(self):
        r = self.client.post(reverse("tasks:task_create"), self.task_payload())
        self.assertEqual(r.status_code, 302)
        self.assertFalse(Task.objects.get(name="Задача с файлами").files.exists())

    def test_edit_task_adds_files_and_keeps_old_ones(self):
        task, _ = self.make_task_with_file("keep.txt")
        self.client.post(
            reverse("tasks:task_edit", kwargs={"pk": task.pk}),
            self.task_payload(name="Задача", attachments=[self.upload("new.txt")]),
        )
        self.assertEqual(
            sorted(task.files.values_list("original_name", flat=True)),
            ["keep.txt", "new.txt"],
        )

    def test_edit_task_without_files_keeps_existing(self):
        task, _ = self.make_task_with_file("keep.txt")
        self.client.post(
            reverse("tasks:task_edit", kwargs={"pk": task.pk}),
            self.task_payload(name="Переименована"),
        )
        self.assertEqual(task.files.count(), 1)

    def test_upload_is_logged_in_history(self):
        self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(attachments=[self.upload("отчёт.txt")]),
        )
        self.assertTrue(
            HistoryEntry.objects.filter(
                owner=self.user,
                text__contains="прикреплён файл «отчёт.txt»",
            ).exists()
        )

    # --- Проверки формы ---

    def test_big_file_accepted(self):
        """Файл в пределах лимита прикрепляется."""
        big = SimpleUploadedFile(
            "big.pdf", b"x" * (3 * 1024 * 1024),
            content_type="application/pdf",
        )
        r = self.client.post(
            reverse("tasks:task_create"), self.task_payload(attachments=big)
        )
        self.assertEqual(r.status_code, 302)
        task = Task.objects.get(name="Задача с файлами")
        self.assertEqual(task.files.count(), 1)

    def test_executable_file_rejected(self):
        r = self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(attachments=self.upload("virus.exe", b"MZ")),
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "прикреплять нельзя")
        self.assertFalse(Task.objects.exists())

    def test_one_bad_file_rejects_whole_upload(self):
        r = self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(
                attachments=[self.upload("ok.txt"), self.upload("bad.bat", b"x")]
            ),
        )
        self.assertEqual(r.status_code, 200)
        self.assertFalse(TaskFile.objects.exists())

    # --- Скачивание и права доступа ---

    def test_owner_can_download_file(self):
        task, task_file = self.make_task_with_file("отчёт.txt")
        r = self.client.get(
            reverse(
                "tasks:task_file_download",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 200)
        self.assertEqual(b"".join(r.streaming_content), b"data")
        self.assertIn("attachment", r["Content-Disposition"])

    def test_anonymous_cannot_download_file(self):
        task, task_file = self.make_task_with_file()
        self.client.logout()
        r = self.client.get(
            reverse(
                "tasks:task_file_download",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 302)
        self.assertIn(reverse("tasks:login"), r["Location"])

    def test_other_user_cannot_download_file(self):
        task, task_file = self.make_task_with_file()
        self.client.force_login(User.objects.create_user("other", password="p"))
        r = self.client.get(
            reverse(
                "tasks:task_file_download",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 404)

    def test_media_is_not_served_directly(self):
        """Вложения не должны быть доступны по прямой ссылке на media/."""
        _, task_file = self.make_task_with_file()
        r = self.client.get(f"{settings.MEDIA_URL}{task_file.file.name}")
        self.assertEqual(r.status_code, 404)

    def test_download_missing_file_returns_404(self):
        task, task_file = self.make_task_with_file()
        os.remove(task_file.file.path)
        r = self.client.get(
            reverse(
                "tasks:task_file_download",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 404)

    # --- Открепление ---

    def test_owner_can_delete_file(self):
        task, task_file = self.make_task_with_file("удалить.txt")
        path = task_file.file.path
        r = self.client.post(
            reverse(
                "tasks:task_file_delete",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertRedirects(r, reverse("tasks:workspace"))
        self.assertFalse(task.files.exists())
        self.assertFalse(os.path.exists(path), "файл остался на диске")
        self.assertTrue(
            HistoryEntry.objects.filter(
                owner=self.user, text__contains="удалён файл «удалить.txt»"
            ).exists()
        )

    def test_other_user_cannot_delete_file(self):
        task, task_file = self.make_task_with_file()
        self.client.force_login(User.objects.create_user("other", password="p"))
        r = self.client.post(
            reverse(
                "tasks:task_file_delete",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 404)
        self.assertTrue(task.files.exists())

    def test_deleting_task_removes_files_from_disk(self):
        task, task_file = self.make_task_with_file()
        path = task_file.file.path
        task.delete()
        self.assertFalse(TaskFile.objects.exists())
        self.assertFalse(os.path.exists(path), "файл остался на диске")

    # --- Страница задачи (удалена) ---

    def test_task_form_has_multiple_file_input(self):
        r = self.client.get(reverse("tasks:task_create"))
        self.assertContains(r, "multiple")
        self.assertContains(r, 'name="attachments"')


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="analysis-test-"))
class AttachmentAnalysisTest(TestCase):
    """Умный разбор вложений: даты и пункты списков из файла."""

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.future = timezone.localdate() + timezone.timedelta(days=30)

    def task_payload(self, **extra):
        payload = {
            "name": "Задача с файлами",
            "description": "Описание",
            "project": "",
            "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
            "priority": Task.Priority.LOW,
            "difficulty": Task.Difficulty.MEDIUM,
            "estimated_duration": Task.EstimatedDuration.UP_TO_30,
            "recurrence": Task.Recurrence.NONE,
            "recurrence_interval_days": "",
        }
        payload.update(extra)
        return payload

    def upload_txt(self, text, name="план.txt"):
        return SimpleUploadedFile(
            name, text.encode("utf-8"), content_type="text/plain"
        )

    def create_task_with_txt(self, text, **extra):
        # Задачи создаются без дедлайна: применение дедлайна из файла
        # тестируется отдельно, поэтому сбрасываем его после создания.
        self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(attachments=self.upload_txt(text), **extra),
        )
        task = Task.objects.get(name="Задача с файлами")
        task.deadline = None
        task.save()
        return task, task.files.first()

    # --- Разбор текста (сервис) ---

    def test_finds_future_numeric_and_word_dates(self):
        text = (
            f"Сдать до {self.future:%d.%m.%Y}. "
            f"Черновик — {self.future.day} "
            f"{list(MONTHS_RU)[self.future.month - 1]} {self.future.year}."
        )
        self.assertEqual(find_dates(text), [self.future.isoformat()])

    def test_past_date_with_year_is_ignored(self):
        self.assertEqual(find_dates("Отчёт за 15.03.2020"), [])

    def test_date_without_year_rolls_to_future(self):
        yesterday = timezone.localdate() - timezone.timedelta(days=1)
        dates = find_dates(f"каждый год {yesterday:%d.%m} праздник")
        expected = yesterday.replace(year=yesterday.year + 1)
        self.assertEqual(dates, [expected.isoformat()])

    def test_checkboxes_preferred_over_bullets(self):
        text = (
            "- просто маркер\n"
            "- [ ] сделать А\n"
            "- [x] уже сделано\n"
            "- [ ] сделать Б\n"
        )
        self.assertEqual(find_items(text), ["сделать А", "сделать Б"])

    def test_bullets_used_when_no_checkboxes(self):
        text = "1. Позвонить\n2) Написать\n- Отправить\n• Забрать\n"
        self.assertEqual(
            find_items(text),
            ["Позвонить", "Написать", "Отправить", "Забрать"],
        )

    # --- Загрузка файлов ---

    def test_upload_stores_analysis(self):
        task, task_file = self.create_task_with_txt(
            f"Дедлайн: {self.future:%d.%m.%Y}\n- [ ] Собрать данные\n"
        )
        self.assertEqual(
            task_file.analysis,
            {"dates": [self.future.isoformat()], "items": ["Собрать данные"]},
        )
        self.assertIsNotNone(task_file.analysis)

    def test_unsupported_file_has_no_analysis(self):
        self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(
                attachments=SimpleUploadedFile("image.png", b"\x89PNG\r\n\x1a\n")
            ),
        )
        task_file = TaskFile.objects.get()
        self.assertIsNone(task_file.analysis)

    def test_corrupt_pdf_does_not_break_upload(self):
        r = self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(
                attachments=SimpleUploadedFile("битый.pdf", b"not a pdf")
            ),
        )
        self.assertEqual(r.status_code, 302)
        self.assertIsNone(TaskFile.objects.get().analysis)

    def test_docx_is_parsed(self):
        import io

        import docx

        document = docx.Document()
        document.add_paragraph("- [ ] Пункт из ворда")
        document.add_paragraph(f"срок {self.future:%d.%m.%Y}")
        buffer = io.BytesIO()
        document.save(buffer)
        self.client.post(
            reverse("tasks:task_create"),
            self.task_payload(
                attachments=SimpleUploadedFile("доклад.docx", buffer.getvalue())
            ),
        )
        analysis = TaskFile.objects.get().analysis
        self.assertEqual(analysis["items"], ["Пункт из ворда"])
        self.assertEqual(analysis["dates"], [self.future.isoformat()])

    # --- Применение дедлайна ---

    def apply_deadline_url(self, task_file):
        return reverse(
            "tasks:task_file_apply_deadline",
            kwargs={"pk": task_file.task_id, "file_pk": task_file.pk},
        )

    def test_apply_deadline_sets_task_deadline(self):
        task, task_file = self.create_task_with_txt(
            f"сдать {self.future:%d.%m.%Y}"
        )
        r = self.client.post(
            self.apply_deadline_url(task_file),
            {"date": self.future.isoformat()},
        )
        self.assertRedirects(r, reverse("tasks:workspace"))
        task.refresh_from_db()
        self.assertEqual(task.deadline, self.future)
        self.assertTrue(
            HistoryEntry.objects.filter(
                owner=self.user, text__contains="из файла «план.txt»"
            ).exists()
        )

    def test_apply_deadline_rejects_date_not_from_file(self):
        task, task_file = self.create_task_with_txt(
            f"сдать {self.future:%d.%m.%Y}"
        )
        other_date = self.future + timezone.timedelta(days=5)
        r = self.client.post(
            self.apply_deadline_url(task_file), {"date": other_date.isoformat()}
        )
        self.assertEqual(r.status_code, 404)
        task.refresh_from_db()
        self.assertIsNone(task.deadline)

    def test_apply_deadline_respects_project_deadline(self):
        project = Project.objects.create(
            owner=self.user,
            name="Proj",
            deadline=self.future - timezone.timedelta(days=10),
        )
        task, task_file = self.create_task_with_txt(
            f"сдать {self.future:%d.%m.%Y}", project=project.pk
        )
        self.client.post(
            self.apply_deadline_url(task_file),
            {"date": self.future.isoformat()},
        )
        task.refresh_from_db()
        self.assertIsNone(task.deadline)  # дедлайн проекта раньше — отказ

    def test_other_user_cannot_apply_deadline(self):
        task, task_file = self.create_task_with_txt(
            f"сдать {self.future:%d.%m.%Y}"
        )
        self.client.force_login(User.objects.create_user("other", password="p"))
        r = self.client.post(
            self.apply_deadline_url(task_file),
            {"date": self.future.isoformat()},
        )
        self.assertEqual(r.status_code, 404)

    # --- Создание задач из пунктов ---

    def create_tasks_url(self, task_file):
        return reverse(
            "tasks:task_file_create_tasks",
            kwargs={"pk": task_file.task_id, "file_pk": task_file.pk},
        )

    def test_create_tasks_from_selected_items(self):
        project = Project.objects.create(owner=self.user, name="Proj")
        task, task_file = self.create_task_with_txt(
            "- [ ] Первый\n- [ ] Второй\n- [ ] Третий\n",
            project=project.pk,
        )
        self.client.post(
            self.create_tasks_url(task_file), {"items": ["0", "2"]}
        )
        created = Task.objects.filter(project=project).exclude(pk=task.pk)
        self.assertEqual(
            sorted(created.values_list("name", flat=True)),
            ["Первый", "Третий"],
        )
        # Использованные пункты исчезли из подсказок, невыбранный остался.
        task_file.refresh_from_db()
        self.assertEqual(task_file.suggested_items, ["Второй"])

    def test_create_tasks_without_selection_creates_nothing(self):
        task, task_file = self.create_task_with_txt("- [ ] Один\n")
        self.client.post(self.create_tasks_url(task_file), {"items": []})
        self.assertEqual(Task.objects.count(), 1)
        task_file.refresh_from_db()
        self.assertEqual(task_file.suggested_items, ["Один"])

    def test_create_tasks_ignores_bad_indices(self):
        task, task_file = self.create_task_with_txt("- [ ] Один\n")
        self.client.post(
            self.create_tasks_url(task_file), {"items": ["5", "abc", "-1"]}
        )
        self.assertEqual(Task.objects.count(), 1)

    def test_other_user_cannot_create_tasks(self):
        task, task_file = self.create_task_with_txt("- [ ] Один\n")
        self.client.force_login(User.objects.create_user("other", password="p"))
        r = self.client.post(self.create_tasks_url(task_file), {"items": ["0"]})
        self.assertEqual(r.status_code, 404)


class GamificationTest(TestCase):
    """Очки, уровни, серии, достижения."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.today = timezone.localdate()

    def complete_task(self, **kwargs):
        """Выполненная задача: по умолчанию — сегодня, без дедлайна."""
        kwargs.setdefault("owner", self.user)
        kwargs.setdefault("name", "T")
        kwargs.setdefault("status", Task.Status.DONE)
        kwargs.setdefault("completed_at", timezone.now())
        return Task.objects.create(**kwargs)

    # --- Момент выполнения ---

    def test_toggle_sets_and_clears_completed_at(self):
        task = Task.objects.create(owner=self.user, name="T")
        self.client.post(reverse("tasks:task_toggle", kwargs={"pk": task.pk}))
        task.refresh_from_db()
        self.assertIsNotNone(task.completed_at)
        self.client.post(reverse("tasks:task_toggle", kwargs={"pk": task.pk}))
        task.refresh_from_db()
        self.assertIsNone(task.completed_at)

    def test_project_complete_stamps_auto_closed_tasks(self):
        project = Project.objects.create(owner=self.user, name="P")
        task = Task.objects.create(owner=self.user, name="T", project=project)
        self.client.post(
            reverse("tasks:project_complete", kwargs={"pk": project.pk}),
            {"action": "complete"},
        )
        task.refresh_from_db()
        self.assertIsNotNone(task.completed_at)

    # --- Очки ---

    def test_task_points_by_priority(self):
        now = timezone.now()
        self.assertEqual(task_points(Task.Priority.LOW, None, now), 10)
        self.assertEqual(task_points(Task.Priority.MEDIUM, None, now), 15)
        self.assertEqual(task_points(Task.Priority.HIGH, None, now), 20)

    def test_task_points_on_time_and_late(self):
        now = timezone.now()
        tomorrow = self.today + timezone.timedelta(days=1)
        yesterday = self.today - timezone.timedelta(days=1)
        self.assertEqual(task_points(Task.Priority.LOW, tomorrow, now), 15)
        self.assertEqual(task_points(Task.Priority.LOW, yesterday, now), 5)

    def test_summary_counts_tasks_and_projects(self):
        self.complete_task()  # 10 очков
        self.complete_task(priority=Task.Priority.HIGH)  # 20 очков
        Project.objects.create(
            owner=self.user, name="P", status=Project.Status.COMPLETED
        )  # 40 очков
        data = summary(self.user)
        self.assertEqual(data["points"], 70)

    def test_reopened_task_loses_points(self):
        task = self.complete_task()
        self.assertEqual(summary(self.user)["points"], 10)
        self.client.post(reverse("tasks:task_toggle", kwargs={"pk": task.pk}))
        self.assertEqual(summary(self.user)["points"], 0)

    def test_other_users_tasks_not_counted(self):
        other = User.objects.create_user("other", password="p")
        Task.objects.create(
            owner=other, name="T", status=Task.Status.DONE,
            completed_at=timezone.now(),
        )
        self.assertEqual(summary(self.user)["points"], 0)

    # --- Уровни ---

    def test_level_thresholds(self):
        self.assertEqual(level_info(0)["name"], "Новичок")
        self.assertEqual(level_info(99)["name"], "Новичок")
        self.assertEqual(level_info(100)["name"], "Стажёр")
        top = level_info(999999)
        self.assertEqual(top["name"], "Разрушитель целей")
        self.assertTrue(top["is_max"])
        self.assertEqual(top["progress"], 100)

    def test_level_progress(self):
        info = level_info(50)  # между 0 и 100
        self.assertEqual(info["progress"], 50)
        self.assertEqual(info["to_next"], 50)
        self.assertEqual(info["next_name"], "Стажёр")

    # --- Серии ---

    def test_streak_counts_consecutive_days(self):
        days = {self.today - timezone.timedelta(days=d) for d in (0, 1, 2)}
        info = streak_info(days, self.today)
        self.assertEqual(info["current"], 3)
        self.assertEqual(info["best"], 3)
        self.assertTrue(info["active_today"])

    def test_streak_alive_until_end_of_day(self):
        # Вчера и позавчера выполнял, сегодня ещё нет — серия жива.
        days = {self.today - timezone.timedelta(days=d) for d in (1, 2)}
        info = streak_info(days, self.today)
        self.assertEqual(info["current"], 2)
        self.assertFalse(info["active_today"])

    def test_streak_broken_by_gap(self):
        days = {
            self.today,
            self.today - timezone.timedelta(days=2),
            self.today - timezone.timedelta(days=3),
            self.today - timezone.timedelta(days=4),
        }
        info = streak_info(days, self.today)
        self.assertEqual(info["current"], 1)
        self.assertEqual(info["best"], 3)

    def test_streak_empty(self):
        info = streak_info(set(), self.today)
        self.assertEqual(info, {"current": 0, "best": 0, "active_today": False})

    # --- Достижения ---

    def achievement(self, title):
        data = summary(self.user)
        return next(
            a for a in achievements(data) if a["title"] == title
        )

    def test_first_task_achievement(self):
        self.assertFalse(self.achievement("Первая задача")["earned"])
        self.complete_task()
        self.assertTrue(self.achievement("Первая задача")["earned"])

    def test_productive_day_achievement_with_progress(self):
        for _ in range(5):
            self.complete_task()
        badge = self.achievement("Продуктивный день")
        self.assertTrue(badge["earned"])
        self.assertEqual(badge["progress"], "5/5")

    def test_finisher_achievement(self):
        self.assertFalse(self.achievement("Финишер")["earned"])
        Project.objects.create(
            owner=self.user, name="P", status=Project.Status.COMPLETED
        )
        self.assertTrue(self.achievement("Финишер")["earned"])

    def test_on_time_achievement_progress(self):
        deadline = self.today + timezone.timedelta(days=1)
        for _ in range(3):
            self.complete_task(deadline=deadline)
        badge = self.achievement("Точно в срок")
        self.assertFalse(badge["earned"])
        self.assertEqual(badge["progress"], "3/10")


class StatsPageTest(TestCase):
    """Страница статистики и элементы геймификации в интерфейсе."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_stats_requires_login(self):
        self.client.logout()
        r = self.client.get(reverse("tasks:stats"))
        self.assertNotEqual(r.status_code, 200)

    def test_stats_page_shows_heatmap_and_achievements(self):
        Task.objects.create(
            owner=self.user, name="T", status=Task.Status.DONE,
            completed_at=timezone.now(),
        )
        r = self.client.get(reverse("tasks:stats"))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Год активности")
        self.assertContains(r, "Выполнено по неделям")
        self.assertContains(r, "Достижения")
        self.assertContains(r, "Первая задача")
        self.assertContains(r, "heatmap__cell--1")  # сегодняшняя клетка окрашена

    def test_stats_page_empty_state(self):
        r = self.client.get(reverse("tasks:stats"))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Всего выполнено")

    def test_no_gamification_ui(self):
        """Серия и звания убраны: ни бейджа в шапке, ни карточки уровня."""
        Task.objects.create(
            owner=self.user, name="T", status=Task.Status.DONE,
            completed_at=timezone.now(),
        )
        for url in (reverse("tasks:workspace"), reverse("tasks:home"),
                    reverse("tasks:stats")):
            with self.subTest(url=url):
                r = self.client.get(url)
                self.assertNotContains(r, "gami-badge")
                self.assertNotContains(r, "Уровень 1")
                self.assertNotContains(r, "Серия")
                self.assertNotContains(r, "Новичок")

    def test_chart_labels_show_week_numbers(self):
        r = self.client.get(reverse("tasks:stats"))
        self.assertContains(r, "нед. ")  # подписи столбцов — номера недель
        self.assertContains(r, ">эта<")  # текущая неделя подписана отдельно

    def test_pages_have_descriptions(self):
        """На страницах есть подсказка для нового пользователя."""
        pages = {
            reverse("tasks:home"): "Сводка на сегодня",
            reverse("tasks:workspace"): "Создавайте задачи",
            reverse("tasks:stats"): "продуктивность в цифрах",
            reverse("tasks:project_list"): "Объединяйте связанные задачи",
            reverse("tasks:note_list"): "Быстрые записи",
            reverse("tasks:history"): "Журнал всех действий",
        }
        for url, text in pages.items():
            with self.subTest(url=url):
                self.assertContains(self.client.get(url), text)

    def test_on_time_share_on_stats(self):
        today = timezone.localdate()
        Task.objects.create(
            owner=self.user, name="OK", status=Task.Status.DONE,
            completed_at=timezone.now(),
            deadline=today + timezone.timedelta(days=1),
        )
        Task.objects.create(
            owner=self.user, name="Late", status=Task.Status.DONE,
            completed_at=timezone.now(),
            deadline=today - timezone.timedelta(days=5),
        )
        r = self.client.get(reverse("tasks:stats"))
        self.assertContains(r, "50%")  # одна из двух — в срок


class AuthSecurityTest(TestCase):
    """Регистрация и вход: дубликаты почты, is_active."""

    def test_register_duplicate_email_rejected(self):
        User.objects.create_user("u", email="user@example.com", password="p")
        r = self.client.post(reverse("tasks:register"), {
            "email": "user@example.com",
            "password1": "Str0ng-pass-123",
            "password2": "Str0ng-pass-123",
        })
        self.assertEqual(r.status_code, 200)  # форма вернулась с ошибкой
        self.assertContains(r, "уже зарегистрирован")
        self.assertEqual(User.objects.count(), 1)

    def test_register_logs_in_new_user(self):
        # Успешная регистрация не должна падать: с несколькими
        # AUTHENTICATION_BACKENDS нужно указывать backend в login().
        r = self.client.post(reverse("tasks:register"), {
            "email": "fresh@example.com",
            "password1": "Str0ng-pass-123",
            "password2": "Str0ng-pass-123",
        })
        self.assertEqual(r.status_code, 302)
        self.assertRedirects(r, reverse("tasks:home"))
        self.assertEqual(User.objects.count(), 1)
        user = User.objects.get(email="fresh@example.com")
        self.assertIn("_auth_user_id", self.client.session)
        self.assertEqual(str(self.client.session["_auth_user_id"]), str(user.pk))

    def test_register_duplicate_email_case_insensitive(self):
        User.objects.create_user("u", email="user@example.com", password="p")
        r = self.client.post(reverse("tasks:register"), {
            "email": "USER@example.com",
            "password1": "Str0ng-pass-123",
            "password2": "Str0ng-pass-123",
        })
        self.assertEqual(User.objects.count(), 1)

    def test_login_survives_legacy_duplicate_emails(self):
        # Старые данные: два аккаунта с одной почтой. Вход не должен
        # падать с 500 — пускаем того, чей пароль подошёл.
        User.objects.create_user("a", email="dup@example.com", password="pass-a")
        User.objects.create_user("b", email="dup@example.com", password="pass-b")
        r = self.client.post(reverse("tasks:login"), {
            "username": "dup@example.com",
            "password": "pass-b",
        })
        self.assertEqual(r.status_code, 302)  # вход удался

    def test_inactive_user_cannot_login_by_email(self):
        user = User.objects.create_user(
            "u", email="user@example.com", password="pass-1"
        )
        user.is_active = False
        user.save()
        r = self.client.post(reverse("tasks:login"), {
            "username": "user@example.com",
            "password": "pass-1",
        })
        self.assertEqual(r.status_code, 200)  # остались на форме входа

    def test_login_email_case_insensitive(self):
        User.objects.create_user("u", email="user@example.com", password="pass-1")
        r = self.client.post(reverse("tasks:login"), {
            "username": "USER@example.com",
            "password": "pass-1",
        })
        self.assertEqual(r.status_code, 302)


class ImportRobustnessTest(TestCase):
    """Импорт: кодировки, дедлайн в сессии, ограничения файла."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.project = Project.objects.create(owner=self.user, name="Proj")
        self.url = reverse("tasks:import_tasks", kwargs={"pk": self.project.pk})

    def import_file(self, name, content, **extra):
        payload = {
            "file": SimpleUploadedFile(name, content, content_type="text/plain"),
            "default_priority": Task.Priority.LOW,
            "default_deadline": "",
        }
        payload.update(extra)
        return self.client.post(self.url, payload)

    def test_cp1251_file_imported(self):
        # Раньше файл читался дважды: второй read() возвращал пустоту,
        # и все cp1251-файлы «не содержали задач».
        r = self.import_file("tasks.txt", "Позвонить в бухгалтерию".encode("cp1251"))
        self.assertContains(r, "Позвонить в бухгалтерию")
        self.client.post(self.url, {"confirm": "1"})
        self.assertTrue(
            Task.objects.filter(name="Позвонить в бухгалтерию").exists()
        )

    def test_import_with_deadline_does_not_crash(self):
        # Раньше datetime.date попадал в JSON-сессию → 500.
        deadline = timezone.localdate() + timezone.timedelta(days=5)
        r = self.import_file(
            "tasks.txt", "Задача".encode("utf-8"),
            default_deadline=deadline.isoformat(),
        )
        self.assertEqual(r.status_code, 200)
        self.client.post(self.url, {"confirm": "1"})
        task = Task.objects.get(name="Задача")
        self.assertEqual(task.deadline, deadline)
        self.assertEqual(task.priority, Task.Priority.LOW)

    def test_import_preview_shows_priority_label(self):
        r = self.import_file(
            "tasks.txt", b"Task", default_priority=Task.Priority.HIGH
        )
        self.assertContains(r, "Приоритет: высокий")

    def test_import_rejects_non_text_extension(self):
        r = self.import_file("tasks.exe", b"MZ data")
        self.assertContains(r, "Поддерживаются только текстовые файлы")
        self.assertFalse(Task.objects.exists())

    def test_import_rejects_huge_file(self):
        r = self.import_file("tasks.txt", b"x" * (2 * 1024 * 1024 + 1))
        self.assertContains(r, "больше 2 МБ")

    def test_import_skips_empty_names_and_truncates_long(self):
        content = "# \n\n- [ ] \n" + "и" * 300 + "\n"
        self.import_file("tasks.md", content.encode("utf-8"))
        self.client.post(self.url, {"confirm": "1"})
        tasks = Task.objects.filter(project=self.project)
        self.assertEqual(tasks.count(), 1)  # пустые имена пропущены
        self.assertEqual(len(tasks.get().name), 200)

    def test_binary_garbage_does_not_crash(self):
        r = self.import_file("tasks.txt", bytes(range(256)) * 4)
        self.assertEqual(r.status_code, 200)


class RedirectSafetyTest(TestCase):
    """Поле next не должно уводить на внешние сайты (open redirect)."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.task = Task.objects.create(owner=self.user, name="T")

    def test_toggle_ignores_external_next(self):
        r = self.client.post(
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk}),
            {"next": "https://evil.example.com/phish"},
        )
        self.assertEqual(r["Location"], reverse("tasks:workspace"))

    def test_toggle_keeps_internal_next(self):
        target = reverse("tasks:overdue_tasks")
        r = self.client.post(
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk}),
            {"next": target},
        )
        self.assertEqual(r["Location"], target)

    def test_theme_toggle_ignores_external_next(self):
        r = self.client.post(
            reverse("tasks:theme_toggle"),
            {"next": "//evil.example.com"},
        )
        self.assertEqual(r["Location"], reverse("tasks:home"))

    def test_history_invalid_week_redirects(self):
        r = self.client.get(reverse("tasks:history") + "?year=2026&week=99")
        self.assertRedirects(r, reverse("tasks:history"))
        r = self.client.get(reverse("tasks:history") + "?year=2026&week=0")
        self.assertRedirects(r, reverse("tasks:history"))


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="ext-test-"))
class AttachmentExtensionBypassTest(TestCase):
    """Обход фильтра расширений хвостовыми точками и пробелами."""

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def try_upload(self, filename):
        return self.client.post(reverse("tasks:task_create"), {
            "name": "T", "description": "Описание", "project": "",
            "deadline": (timezone.localdate() + timezone.timedelta(days=7)).isoformat(),
            "priority": Task.Priority.LOW,
            "difficulty": Task.Difficulty.MEDIUM,
            "estimated_duration": Task.EstimatedDuration.UP_TO_30,
            "recurrence": Task.Recurrence.NONE,
            "recurrence_interval_days": "",
            "attachments": SimpleUploadedFile(filename, b"MZ"),
        })

    def test_trailing_dot_rejected(self):
        self.try_upload("virus.exe.")
        self.assertFalse(TaskFile.objects.exists())

    def test_uppercase_rejected(self):
        self.try_upload("VIRUS.EXE")
        self.assertFalse(TaskFile.objects.exists())


class TemplateRunTest(TestCase):
    """Запуск процесса по шаблону: выбор шагов, история запусков."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.template = ProjectTemplate.objects.create(
            owner=self.user, name="Онбординг", description="Новый сотрудник"
        )
        self.s1 = TemplateTask.objects.create(
            template=self.template, name="Завести почту", deadline_offset_days=1
        )
        self.s2 = TemplateTask.objects.create(
            template=self.template, name="Выдать ноутбук"
        )
        self.s3 = TemplateTask.objects.create(
            template=self.template, name="Экскурсия по офису"
        )
        self.url = reverse("tasks:template_run", kwargs={"pk": self.template.pk})

    def test_run_page_lists_steps(self):
        r = self.client.get(self.url)
        self.assertContains(r, "Завести почту")
        self.assertContains(r, "Запустить процесс")

    def test_run_copies_selected_steps_only(self):
        r = self.client.post(self.url, {
            "name": "Онбординг: Мария",
            "deadline": "",
            "steps": [self.s1.pk, self.s2.pk],
        })
        project = Project.objects.get()
        self.assertEqual(r["Location"],
                         reverse("tasks:project_detail", kwargs={"pk": project.pk}))
        self.assertEqual(project.name, "Онбординг: Мария")
        self.assertEqual(project.source_template, self.template)
        names = set(project.tasks.values_list("name", flat=True))
        self.assertEqual(names, {"Завести почту", "Выдать ноутбук"})

    def test_run_without_steps_rejected(self):
        r = self.client.post(self.url, {"name": "", "deadline": "", "steps": []})
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "хотя бы один шаг")
        self.assertFalse(Project.objects.exists())

    def test_run_default_name_from_template(self):
        self.client.post(self.url, {
            "name": "", "deadline": "", "steps": [self.s1.pk],
        })
        self.assertEqual(Project.objects.get().name, "Онбординг")

    def test_foreign_template_404(self):
        other = User.objects.create_user("v", password="p")
        foreign = ProjectTemplate.objects.create(owner=other, name="Чужой")
        r = self.client.get(
            reverse("tasks:template_run", kwargs={"pk": foreign.pk})
        )
        self.assertEqual(r.status_code, 404)

    def test_foreign_steps_ignored(self):
        # Шаг из чужого шаблона нельзя протащить в свой запуск.
        other = User.objects.create_user("v", password="p")
        foreign = ProjectTemplate.objects.create(owner=other, name="Чужой")
        foreign_step = TemplateTask.objects.create(
            template=foreign, name="Чужой шаг"
        )
        r = self.client.post(self.url, {
            "name": "", "deadline": "", "steps": [foreign_step.pk],
        })
        self.assertEqual(r.status_code, 200)  # форма с ошибкой
        self.assertFalse(Project.objects.exists())

    def test_template_detail_shows_runs_with_progress(self):
        self.client.post(self.url, {
            "name": "Запуск 1", "deadline": "", "steps": [self.s1.pk, self.s2.pk],
        })
        project = Project.objects.get()
        task = project.tasks.first()
        task.status = Task.Status.DONE
        task.completed_at = timezone.now()
        task.save()
        r = self.client.get(
            reverse("tasks:template_detail", kwargs={"pk": self.template.pk})
        )
        self.assertContains(r, "Запуск 1")
        self.assertContains(r, "1/2")

    def test_project_list_shows_progress_bar(self):
        project = Project.objects.create(owner=self.user, name="P")
        Task.objects.create(
            owner=self.user, project=project, name="A",
            status=Task.Status.DONE, completed_at=timezone.now(),
        )
        Task.objects.create(owner=self.user, project=project, name="B")
        r = self.client.get(reverse("tasks:project_list"))
        self.assertContains(r, "pbar")
        self.assertContains(r, "width: 50%")


class TaskDetailViewTest(TestCase):
    """Просмотр задачи: доступ, данные слева, файлы и слайды справа."""

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.other = User.objects.create_user("v", password="p")
        self.client.force_login(self.user)

    def make_task(self, **extra):
        task = Task.objects.create(owner=self.user, name="Задача", **extra)
        return task

    def make_file(self, task, name="doc.txt", content=b"data"):
        return TaskFile.objects.create(
            task=task,
            file=SimpleUploadedFile(name, content, content_type="text/plain"),
            original_name=name,
        )

    def detail_url(self, task):
        return reverse("tasks:task_detail", kwargs={"pk": task.pk})

    def test_owner_can_open_task_page(self):
        task = self.make_task(description="Описание задачи")
        r = self.client.get(self.detail_url(task))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Задача")
        self.assertContains(r, "Описание задачи")

    def test_foreign_user_gets_404(self):
        task = self.make_task()
        self.client.force_login(self.other)
        r = self.client.get(self.detail_url(task))
        self.assertEqual(r.status_code, 404)

    def test_anonymous_redirected_to_login(self):
        self.client.logout()
        task = self.make_task()
        r = self.client.get(self.detail_url(task))
        self.assertEqual(r.status_code, 302)

    def test_page_shows_files_and_download_button(self):
        task = self.make_task()
        self.make_file(task, name="отчёт.txt", content=b"hello")
        r = self.client.get(self.detail_url(task))
        self.assertContains(r, "отчёт.txt")
        self.assertContains(r, "Скачать")

    def test_text_file_split_into_slides_by_separator(self):
        task = self.make_task()
        content = "Слайд один\n\n---\n\nСлайд два\n\n---\n\nСлайд три".encode()
        self.make_file(task, name="презентация.md", content=content)
        r = self.client.get(self.detail_url(task))
        self.assertContains(r, "Слайд один")
        self.assertContains(r, "Слайд два")
        self.assertContains(r, "Слайд три")
        self.assertContains(r, "data-slide=", count=3)
        self.assertContains(r, "data-current")

    def test_text_file_without_separator_is_one_slide(self):
        task = self.make_task()
        self.make_file(task, name="заметки.txt", content="один\nдва\nтри".encode())
        r = self.client.get(self.detail_url(task))
        self.assertContains(r, "один")
        self.assertContains(r, "три")
        self.assertNotContains(r, "file-view__counter")

    def test_unsupported_file_shows_empty_body(self):
        task = self.make_task()
        self.make_file(task, name="вирус.exe", content=b"MZ")
        r = self.client.get(self.detail_url(task))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "вирус.exe")

    def test_image_file_uses_preview_url(self):
        task = self.make_task()
        self.make_file(task, name="картинка.png", content=b"\x89PNG")
        r = self.client.get(self.detail_url(task))
        self.assertContains(r, f"/tasks/{task.pk}/files/1/preview/")
        self.assertContains(r, "file-view__image")

    def test_preview_endpoint_serves_image_inline(self):
        task = self.make_task()
        task_file = self.make_file(task, name="картинка.png", content=b"\x89PNGdata")
        r = self.client.get(
            reverse(
                "tasks:task_file_preview",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r["Content-Type"], "image/png")
        self.assertNotIn("attachment", r["Content-Disposition"])
        self.assertEqual(b"".join(r.streaming_content), b"\x89PNGdata")

    def test_preview_endpoint_blocks_non_images(self):
        task = self.make_task()
        task_file = self.make_file(task, name="заметки.txt")
        r = self.client.get(
            reverse(
                "tasks:task_file_preview",
                kwargs={"pk": task.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 404)

    def test_edit_link_goes_to_detail(self):
        task = self.make_task()
        r = self.client.get(self.detail_url(task))
        self.assertContains(r, "Изменить")


class ProjectFileTest(TestCase):
    """Файлы проекта: прикрепление в форме, просмотр, действия с файлами."""

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.other = User.objects.create_user("v", password="p")
        self.client.force_login(self.user)

    def upload(self, name, content=b"data"):
        return SimpleUploadedFile(name, content, content_type="text/plain")

    def make_project(self, **extra):
        defaults = {
            "name": "Проект",
            "description": "Описание",
            "deadline": timezone.localdate() + timezone.timedelta(days=30),
        }
        defaults.update(extra)
        return Project.objects.create(owner=self.user, **defaults)

    def project_payload(self, **extra):
        """Минимальный валидный POST для формы проекта (без задач)."""
        payload = {
            "name": "Проект",
            "description": "Описание",
            "deadline": (
                timezone.localdate() + timezone.timedelta(days=30)
            ).isoformat(),
            "tasks-TOTAL_FORMS": "0",
            "tasks-INITIAL_FORMS": "0",
            "tasks-MIN_NUM_FORMS": "0",
            "tasks-MAX_NUM_FORMS": "1000",
        }
        payload.update(extra)
        return payload

    # --- Прикрепление в форме проекта ---

    def test_create_project_attaches_file_to_project(self):
        r = self.client.post(
            reverse("tasks:project_create"),
            self.project_payload(
                name="Проект с файлом",
                attachments=[self.upload("doc.txt")],
                attach_target=[""],
            ),
        )
        self.assertEqual(r.status_code, 302)
        project = Project.objects.get(name="Проект с файлом")
        task_file = project.files.get()
        self.assertIsNone(task_file.task)
        self.assertEqual(task_file.project, project)
        self.assertEqual(task_file.original_name, "doc.txt")

    def test_edit_project_binds_file_to_task(self):
        project = self.make_project()
        task = Task.objects.create(
            owner=self.user, project=project, name="Задача проекта"
        )
        payload = self.project_payload(
            attachments=[self.upload("doc.txt")],
            attach_target=[str(task.pk)],
            **{
                "tasks-TOTAL_FORMS": "1",
                "tasks-INITIAL_FORMS": "1",
                "tasks-0-id": str(task.pk),
                "tasks-0-name": "Задача проекта",
                "tasks-0-deadline": "",
                "tasks-0-description": "",
                "tasks-0-priority": Task.Priority.LOW,
                "tasks-0-difficulty": Task.Difficulty.MEDIUM,
                "tasks-0-estimated_duration": Task.EstimatedDuration.UP_TO_30,
            },
        )
        r = self.client.post(
            reverse("tasks:project_edit", kwargs={"pk": project.pk}), payload
        )
        self.assertEqual(r.status_code, 302)
        task_file = TaskFile.objects.get()
        self.assertEqual(task_file.task, task)
        self.assertIsNone(task_file.project)

    def test_edit_project_file_without_target_stays_on_project(self):
        project = self.make_project()
        task = Task.objects.create(
            owner=self.user, project=project, name="Задача проекта"
        )
        payload = self.project_payload(
            attachments=[self.upload("doc.txt")],
            attach_target=[""],
            **{
                "tasks-TOTAL_FORMS": "1",
                "tasks-INITIAL_FORMS": "1",
                "tasks-0-id": str(task.pk),
                "tasks-0-name": "Задача проекта",
                "tasks-0-deadline": "",
                "tasks-0-description": "",
                "tasks-0-priority": Task.Priority.LOW,
                "tasks-0-difficulty": Task.Difficulty.MEDIUM,
                "tasks-0-estimated_duration": Task.EstimatedDuration.UP_TO_30,
            },
        )
        self.client.post(
            reverse("tasks:project_edit", kwargs={"pk": project.pk}), payload
        )
        task_file = TaskFile.objects.get()
        self.assertIsNone(task_file.task)
        self.assertEqual(task_file.project, project)

    def test_attach_target_from_foreign_task_ignored(self):
        project = self.make_project()
        foreign = Task.objects.create(owner=self.other, name="Чужой")
        r = self.client.post(
            reverse("tasks:project_create"),
            self.project_payload(
                name="Проект2",
                attachments=[self.upload("doc.txt")],
                attach_target=[str(foreign.pk)],
            ),
        )
        self.assertEqual(r.status_code, 302)
        project2 = Project.objects.get(name="Проект2")
        task_file = project2.files.get()
        self.assertIsNone(task_file.task)
        self.assertEqual(task_file.project, project2)

    # --- Просмотр проекта ---

    def test_project_detail_shows_project_files_tab(self):
        project = self.make_project()
        TaskFile.objects.create(
            project=project,
            file=self.upload("proj.txt"),
            original_name="proj.txt",
        )
        r = self.client.get(
            reverse("tasks:project_detail", kwargs={"pk": project.pk})
        )
        self.assertContains(r, "proj.txt")
        self.assertContains(r, "Файлы проекта")

    def test_project_detail_shows_task_files_tab(self):
        project = self.make_project()
        task = Task.objects.create(
            owner=self.user, project=project, name="Задача с файлом"
        )
        TaskFile.objects.create(
            task=task,
            file=self.upload("task.txt"),
            original_name="task.txt",
        )
        r = self.client.get(
            reverse("tasks:project_detail", kwargs={"pk": project.pk})
        )
        self.assertContains(r, "task.txt")
        self.assertContains(
            r, f'data-file-tab="task-{task.pk}"'
        )

    # --- Действия с файлами проекта ---

    def test_project_file_apply_deadline(self):
        project = self.make_project(deadline=None)
        task_file = TaskFile.objects.create(
            project=project,
            file=self.upload("doc.txt"),
            original_name="doc.txt",
            analysis={"dates": ["2026-12-01"], "items": []},
        )
        r = self.client.post(
            reverse(
                "tasks:project_file_apply_deadline",
                kwargs={"pk": project.pk, "file_pk": task_file.pk},
            ),
            {"date": "2026-12-01"},
        )
        self.assertEqual(r.status_code, 302)
        project.refresh_from_db()
        self.assertEqual(project.deadline, datetime.date(2026, 12, 1))

    def test_project_file_apply_deadline_ignores_foreign_dates(self):
        project = self.make_project(deadline=None)
        task_file = TaskFile.objects.create(
            project=project,
            file=self.upload("doc.txt"),
            original_name="doc.txt",
            analysis={"dates": ["2026-12-01"], "items": []},
        )
        r = self.client.post(
            reverse(
                "tasks:project_file_apply_deadline",
                kwargs={"pk": project.pk, "file_pk": task_file.pk},
            ),
            {"date": "2030-01-01"},
        )
        self.assertEqual(r.status_code, 404)
        project.refresh_from_db()
        self.assertIsNone(project.deadline)

    def test_project_file_create_tasks(self):
        project = self.make_project()
        task_file = TaskFile.objects.create(
            project=project,
            file=self.upload("doc.txt"),
            original_name="doc.txt",
            analysis={"dates": [], "items": ["Пункт 1", "Пункт 2"]},
        )
        r = self.client.post(
            reverse(
                "tasks:project_file_create_tasks",
                kwargs={"pk": project.pk, "file_pk": task_file.pk},
            ),
            {"items": ["0", "1"]},
        )
        self.assertEqual(r.status_code, 302)
        names = list(
            project.tasks.values_list("name", flat=True).order_by("name")
        )
        self.assertEqual(names, ["Пункт 1", "Пункт 2"])

    def test_project_file_delete(self):
        project = self.make_project()
        task_file = TaskFile.objects.create(
            project=project,
            file=self.upload("doc.txt"),
            original_name="doc.txt",
        )
        r = self.client.post(
            reverse(
                "tasks:project_file_delete",
                kwargs={"pk": project.pk, "file_pk": task_file.pk},
            ),
        )
        self.assertEqual(r.status_code, 302)
        self.assertFalse(TaskFile.objects.filter(pk=task_file.pk).exists())

    def test_foreign_user_cannot_download_project_file(self):
        project = self.make_project()
        task_file = TaskFile.objects.create(
            project=project,
            file=self.upload("doc.txt"),
            original_name="doc.txt",
        )
        self.client.force_login(self.other)
        r = self.client.get(
            reverse(
                "tasks:project_file_download",
                kwargs={"pk": project.pk, "file_pk": task_file.pk},
            )
        )
        self.assertEqual(r.status_code, 404)
        r = self.client.post(
            reverse(
                "tasks:project_file_delete",
                kwargs={"pk": project.pk, "file_pk": task_file.pk},
            ),
        )
        self.assertEqual(r.status_code, 404)
        self.assertTrue(TaskFile.objects.filter(pk=task_file.pk).exists())


class TaskToggleConfirmTest(TestCase):
    """Страницы подтверждения завершения / возврата задачи."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.task = Task.objects.create(
            owner=self.user, name="Задача", description="Описание"
        )

    def test_get_confirm_page_for_completing(self):
        r = self.client.get(
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk})
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Завершить задачу?")
        self.assertContains(r, "Задача")

    def test_get_confirm_page_for_returning(self):
        self.task.status = Task.Status.DONE
        self.task.save()
        r = self.client.get(
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk})
        )
        self.assertContains(r, "Вернуть задачу в работу?")

    def test_get_confirm_page_foreign_task_404(self):
        other = User.objects.create_user("v", password="p")
        self.client.force_login(other)
        r = self.client.get(
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk})
        )
        self.assertEqual(r.status_code, 404)

    def test_post_still_toggles(self):
        r = self.client.post(
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk}),
            {"next": reverse("tasks:workspace")},
        )
        self.assertEqual(r.status_code, 302)
        self.assertTrue(Task.objects.get(pk=self.task.pk).is_done)


class ProjectReopenConfirmTest(TestCase):
    """Страница подтверждения возврата завершённого проекта."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def test_get_confirm_page(self):
        project = Project.objects.create(
            owner=self.user,
            name="Проект",
            description="Описание",
            deadline=timezone.localdate(),
            status=Project.Status.COMPLETED,
        )
        r = self.client.get(
            reverse("tasks:project_reopen", kwargs={"pk": project.pk})
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Вернуть проект в работу?")

    def test_get_confirm_page_for_active_project_redirects(self):
        project = Project.objects.create(
            owner=self.user,
            name="Проект",
            description="Описание",
            deadline=timezone.localdate(),
        )
        r = self.client.get(
            reverse("tasks:project_reopen", kwargs={"pk": project.pk})
        )
        self.assertEqual(r.status_code, 302)

    def test_post_reopens(self):
        project = Project.objects.create(
            owner=self.user,
            name="Проект",
            description="Описание",
            deadline=timezone.localdate(),
            status=Project.Status.COMPLETED,
        )
        r = self.client.post(
            reverse("tasks:project_reopen", kwargs={"pk": project.pk})
        )
        self.assertEqual(r.status_code, 302)
        project.refresh_from_db()
        self.assertFalse(project.is_completed)


class AddTaskFromProjectFormTest(TestCase):
    """Кнопка «Создать задачу» в форме проекта."""

    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)

    def payload(self, **extra):
        payload = {
            "name": "Новый проект",
            "description": "Описание",
            "deadline": (
                timezone.localdate() + timezone.timedelta(days=30)
            ).isoformat(),
            "add_task": "1",
            "tasks-TOTAL_FORMS": "0",
            "tasks-INITIAL_FORMS": "0",
            "tasks-MIN_NUM_FORMS": "0",
            "tasks-MAX_NUM_FORMS": "1000",
        }
        payload.update(extra)
        return payload

    def task_row(self, index=0, **fields):
        """Строка формсета задач с заполненным названием."""
        row = {
            "id": "",
            "name": "Задача из строки",
            "description": "Описание задачи",
            "deadline": "",
            "priority": "0",
            "difficulty": "1",
            "estimated_duration": "1",
        }
        row.update(fields)
        return {f"tasks-{index}-{key}": value for key, value in row.items()}

    def test_create_project_via_add_task_button(self):
        r = self.client.post(
            reverse("tasks:project_create"),
            self.payload(**{
                "tasks-TOTAL_FORMS": "1",
                **self.task_row(),
            }),
        )
        project = Project.objects.get(name="Новый проект")
        self.assertRedirects(
            r, reverse("tasks:project_edit", kwargs={"pk": project.pk})
        )
        task = project.tasks.get()
        self.assertEqual(task.name, "Задача из строки")
        self.assertEqual(task.description, "Описание задачи")

    def test_edit_project_adds_task_via_button(self):
        project = Project.objects.create(
            owner=self.user,
            name="Проект",
            description="Описание",
            deadline=timezone.localdate() + timezone.timedelta(days=30),
        )
        r = self.client.post(
            reverse("tasks:project_edit", kwargs={"pk": project.pk}),
            self.payload(
                name="Проект",
                **{
                    "tasks-TOTAL_FORMS": "1",
                    **self.task_row(),
                },
            ),
        )
        self.assertEqual(r.status_code, 302)
        task = project.tasks.get()
        self.assertEqual(task.name, "Задача из строки")

    def test_add_task_takes_last_filled_row(self):
        # Несколько строк: задача создаётся из последней заполненной.
        r = self.client.post(
            reverse("tasks:project_create"),
            self.payload(**{
                "tasks-TOTAL_FORMS": "2",
                **self.task_row(0, name="Первая"),
                **self.task_row(1, name="Последняя"),
            }),
        )
        project = Project.objects.get(name="Новый проект")
        self.assertRedirects(
            r, reverse("tasks:project_edit", kwargs={"pk": project.pk})
        )
        self.assertEqual(project.tasks.count(), 1)
        self.assertEqual(project.tasks.get().name, "Последняя")

    def test_add_task_skips_existing_rows_in_edit(self):
        # При редактировании уже сохранённые задачи не дублируются:
        # создаётся только задача из новой строки.
        project = Project.objects.create(
            owner=self.user,
            name="Проект",
            description="Описание",
            deadline=timezone.localdate() + timezone.timedelta(days=30),
        )
        old_task = Task.objects.create(owner=self.user, project=project, name="Старая")
        r = self.client.post(
            reverse("tasks:project_edit", kwargs={"pk": project.pk}),
            self.payload(
                name="Проект",
                **{
                    "tasks-TOTAL_FORMS": "2",
                    "tasks-INITIAL_FORMS": "1",
                    "tasks-0-id": str(old_task.pk),
                    "tasks-0-name": "Старая",
                    "tasks-0-deadline": "",
                    "tasks-0-description": "",
                    "tasks-0-priority": "0",
                    "tasks-0-difficulty": "1",
                    "tasks-0-estimated_duration": "1",
                    **self.task_row(1),
                },
            ),
        )
        self.assertEqual(r.status_code, 302)
        self.assertEqual(project.tasks.count(), 2)
        self.assertIsNotNone(project.tasks.filter(name="Задача из строки").first())

    def test_add_task_with_empty_last_row_just_saves_project(self):
        r = self.client.post(
            reverse("tasks:project_create"), self.payload()
        )
        project = Project.objects.get(name="Новый проект")
        self.assertRedirects(
            r, reverse("tasks:project_edit", kwargs={"pk": project.pk})
        )
        self.assertEqual(project.tasks.count(), 0)

    def test_add_task_with_empty_name_shows_errors(self):
        r = self.client.post(
            reverse("tasks:project_create"), self.payload(name="")
        )
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Укажите название")


class TaskShareTest(TestCase):
    """Публичные ссылки на шаги: включение, выключение, выполнение."""

    def setUp(self):
        self.user = User.objects.create_user("secret_owner_99", password="p")
        self.client.force_login(self.user)
        self.task = Task.objects.create(
            owner=self.user, name="Подписать договор", description="Секретно"
        )

    def enable(self):
        self.client.post(
            reverse("tasks:task_share_enable", kwargs={"pk": self.task.pk})
        )
        self.task.refresh_from_db()

    def test_enable_creates_code_and_public_page_opens(self):
        self.enable()
        self.assertTrue(self.task.share_code)
        anon = self.client_class()  # другой браузер, без входа
        r = anon.get(
            reverse("tasks:public_task", kwargs={"code": self.task.share_code})
        )
        self.assertContains(r, "Подписать договор")
        self.assertContains(r, "Отметить выполненным")

    def test_reissue_changes_code(self):
        self.enable()
        old = self.task.share_code
        self.enable()
        self.assertNotEqual(self.task.share_code, old)
        anon = self.client_class()
        r = anon.get(reverse("tasks:public_task", kwargs={"code": old}))
        self.assertEqual(r.status_code, 404)

    def test_disable_kills_link(self):
        self.enable()
        code = self.task.share_code
        self.client.post(
            reverse("tasks:task_share_disable", kwargs={"pk": self.task.pk})
        )
        self.task.refresh_from_db()
        self.assertIsNone(self.task.share_code)
        anon = self.client_class()
        r = anon.get(reverse("tasks:public_task", kwargs={"code": code}))
        self.assertEqual(r.status_code, 404)

    def test_anonymous_completes_task(self):
        self.enable()
        anon = self.client_class()
        anon.post(reverse(
            "tasks:public_task_complete", kwargs={"code": self.task.share_code}
        ))
        self.task.refresh_from_db()
        self.assertTrue(self.task.is_done)
        self.assertIsNotNone(self.task.completed_at)
        self.assertTrue(HistoryEntry.objects.filter(
            owner=self.user, text__contains="по публичной ссылке"
        ).exists())

    def test_double_complete_is_idempotent(self):
        self.task.recurrence = Task.Recurrence.DAILY
        self.task.save()
        self.enable()
        anon = self.client_class()
        url = reverse(
            "tasks:public_task_complete", kwargs={"code": self.task.share_code}
        )
        anon.post(url)
        anon.post(url)
        # Повторная задача «возродилась» ровно один раз
        self.assertEqual(
            Task.objects.filter(name=self.task.name).count(), 2
        )

    def test_share_management_requires_owner(self):
        other = User.objects.create_user("v", password="p")
        self.client.force_login(other)
        r = self.client.post(
            reverse("tasks:task_share_enable", kwargs={"pk": self.task.pk})
        )
        self.assertEqual(r.status_code, 404)
        self.task.refresh_from_db()
        self.assertIsNone(self.task.share_code)

    def test_public_page_hides_files_and_owner(self):
        self.enable()
        anon = self.client_class()
        r = anon.get(
            reverse("tasks:public_task", kwargs={"code": self.task.share_code})
        )
        self.assertNotContains(r, "Прикреплённые файлы")
        self.assertNotContains(r, self.user.username)


class JournalTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("u", password="p")
        self.client.force_login(self.user)
        self.project = Project.objects.create(owner=self.user, name="Proj")
        self.today = timezone.localdate()

    def _done_task(self, name, days_ago=0, project=None):
        task = Task.objects.create(
            owner=self.user,
            name=name,
            project=project,
            status=Task.Status.DONE,
        )
        task.completed_at = timezone.now() - timezone.timedelta(days=days_ago)
        task.save(update_fields=["completed_at"])
        return task

    def test_journal_page_loads(self):
        r = self.client.get(reverse("tasks:journal"))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Журнал достижений")
        self.assertContains(r, "Собрать сводку")

    def test_add_entry(self):
        r = self.client.post(
            reverse("tasks:journal"),
            {"text": "Сделал демо", "date": self.today.isoformat(),
             "project": str(self.project.pk)},
        )
        self.assertEqual(r.status_code, 302)
        entry = JournalEntry.objects.get(owner=self.user)
        self.assertEqual(entry.text, "Сделал демо")
        self.assertEqual(entry.project, self.project)

    def test_entries_grouped_by_day_and_reminder(self):
        JournalEntry.objects.create(
            owner=self.user, date=self.today, text="Сегодняшняя"
        )
        yesterday = self.today - timezone.timedelta(days=1)
        JournalEntry.objects.create(
            owner=self.user, date=yesterday, text="Вчерашняя"
        )
        r = self.client.get(reverse("tasks:journal"))
        self.assertContains(r, "Сегодняшняя")
        self.assertContains(r, "Вчерашняя")
        # Запись за сегодня есть — напоминания нет.
        self.assertNotContains(r, "Сегодня ещё нет записи")

    def test_reminder_shown_when_no_entry_today(self):
        r = self.client.get(reverse("tasks:journal"))
        self.assertContains(r, "Сегодня ещё нет записи")
        # И на главной появился баннер.
        r = self.client.get(reverse("tasks:home"))
        self.assertContains(r, "Журнал: сегодня ещё нет записи")

    def test_no_banner_after_entry_today(self):
        JournalEntry.objects.create(
            owner=self.user, date=self.today, text="Уже записал"
        )
        r = self.client.get(reverse("tasks:home"))
        self.assertNotContains(r, "Журнал: сегодня ещё нет записи")

    def test_edit_and_delete_entry(self):
        entry = JournalEntry.objects.create(
            owner=self.user, date=self.today, text="Старая"
        )
        r = self.client.post(
            reverse("tasks:journal_entry_edit", kwargs={"pk": entry.pk}),
            {"text": "Новая", "date": self.today.isoformat(), "project": ""},
        )
        self.assertEqual(r.status_code, 302)
        entry.refresh_from_db()
        self.assertEqual(entry.text, "Новая")
        r = self.client.post(
            reverse("tasks:journal_entry_delete", kwargs={"pk": entry.pk})
        )
        self.assertEqual(r.status_code, 302)
        self.assertFalse(JournalEntry.objects.exists())

    def test_cannot_touch_other_users_entries(self):
        other = User.objects.create_user("v", password="p")
        entry = JournalEntry.objects.create(
            owner=other, date=self.today, text="Чужая"
        )
        for url_name in ("tasks:journal_entry_edit", "tasks:journal_entry_delete"):
            r = self.client.get(reverse(url_name, kwargs={"pk": entry.pk}))
            self.assertEqual(r.status_code, 404)

    def test_summary_collects_entries_and_tasks(self):
        JournalEntry.objects.create(
            owner=self.user, date=self.today, text="Подготовил отчёт",
            project=self.project,
        )
        self._done_task("Сделанное", project=self.project)
        old = self._done_task("Давнее", days_ago=40)
        r = self.client.get(reverse("tasks:journal_summary"))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, "Подготовил отчёт")
        self.assertContains(r, "Сделанное")
        # Задача 40-дневной давности не попадает в сводку недели.
        self.assertNotContains(r, "Давнее")

    def test_summary_periods(self):
        self._done_task("Давнее", days_ago=40)
        r = self.client.get(reverse("tasks:journal_summary") + "?period=halfyear")
        self.assertContains(r, "Давнее")
        r = self.client.get(reverse("tasks:journal_summary") + "?period=month")
        self.assertNotContains(r, "Давнее")

    def test_summary_includes_completed_projects(self):
        self.project.status = Project.Status.COMPLETED
        self.project.completed_at = timezone.now()
        self.project.save(update_fields=["status", "completed_at"])
        r = self.client.get(reverse("tasks:journal_summary"))
        self.assertContains(r, "Завершённые проекты")
        self.assertContains(r, "Proj")

    def test_markdown_export_downloads_file(self):
        JournalEntry.objects.create(
            owner=self.user, date=self.today, text="Запись"
        )
        r = self.client.get(reverse("tasks:journal_summary_md"))
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r["Content-Type"], "text/markdown; charset=utf-8")
        self.assertIn("attachment", r["Content-Disposition"])
        content = r.content.decode("utf-8")
        self.assertIn("# Что сделано", content)
        self.assertIn("Запись", content)

    def test_project_complete_sets_completed_at(self):
        self.client.post(
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            {"action": "complete"},
        )
        self.project.refresh_from_db()
        self.assertIsNotNone(self.project.completed_at)
        # Переоткрытие сбрасывает дату.
        self.client.post(
            reverse("tasks:project_reopen", kwargs={"pk": self.project.pk})
        )
        self.project.refresh_from_db()
        self.assertIsNone(self.project.completed_at)

    def test_entry_streak_counts_consecutive_days(self):
        from tasks.services.journal_service import entry_streak

        self.assertEqual(entry_streak(self.user, self.today), 0)
        JournalEntry.objects.create(
            owner=self.user, date=self.today, text="1"
        )
        JournalEntry.objects.create(
            owner=self.user,
            date=self.today - timezone.timedelta(days=1),
            text="2",
        )
        self.assertEqual(entry_streak(self.user, self.today), 2)
        # Дыра вчера: сегодняшняя запись — серия из одного дня.
        JournalEntry.objects.filter(text="2").delete()
        self.assertEqual(entry_streak(self.user, self.today), 1)


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="fileval-test-"))
class FileValidationTest(TestCase):
    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(settings.MEDIA_ROOT, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user(
            username="filetest", password="pass123"
        )
        self.client.force_login(self.user)

    def test_valid_file_accepted(self):
        resp = self.client.post(
            reverse("tasks:task_create"),
            _task_payload(attachments=[SimpleUploadedFile("notes.txt", b"data", content_type="text/plain")]),
        )
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(TaskFile.objects.exists())

    def test_too_large_file_rejected(self):
        big = b"x" * (settings.MAX_TASK_FILE_SIZE + 1)
        resp = self.client.post(
            reverse("tasks:task_create"),
            _task_payload(attachments=[SimpleUploadedFile("big.txt", big, content_type="text/plain")]),
        )
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(TaskFile.objects.exists())

    def test_blocked_extension_rejected(self):
        resp = self.client.post(
            reverse("tasks:task_create"),
            _task_payload(attachments=[SimpleUploadedFile("virus.exe", b"evil", content_type="application/octet-stream")]),
        )
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(TaskFile.objects.exists())

    def test_max_files_exceeded(self):
        task = Task.objects.create(owner=self.user, name="Test task")
        url = reverse("tasks:task_edit", kwargs={"pk": task.pk})
        files = [
            SimpleUploadedFile(f"f{i}.txt", b"content", content_type="text/plain")
            for i in range(settings.MAX_TASK_FILES_PER_TASK + 1)
        ]
        resp = self.client.post(
            url,
            _task_payload(name="Test task", attachments=files),
        )
        self.assertEqual(resp.status_code, 200)


class SecuritySettingsTest(TestCase):
    def test_x_frame_options_deny(self):
        self.assertEqual(settings.X_FRAME_OPTIONS, "DENY")

    def test_content_type_nosniff(self):
        self.assertTrue(settings.SECURE_CONTENT_TYPE_NOSNIFF)

    def test_max_file_size_setting(self):
        self.assertEqual(settings.MAX_TASK_FILE_SIZE, 10 * 1024 * 1024)

    def test_max_files_per_task_setting(self):
        self.assertEqual(settings.MAX_TASK_FILES_PER_TASK, 10)


class ThrottleTest(TestCase):
    def test_rate_limit_allows_normal_usage(self):
        from tasks.services.throttle import check_rate_limit
        from django.http import HttpRequest

        request = HttpRequest()
        request.session = {}
        allowed, _ = check_rate_limit(request, "test_action", max_actions=3, period_seconds=60)
        self.assertTrue(allowed)

    def test_rate_limit_blocks_excess(self):
        from tasks.services.throttle import check_rate_limit
        from django.http import HttpRequest

        request = HttpRequest()
        request.session = {}
        for i in range(3):
            allowed, _ = check_rate_limit(request, "test_action", max_actions=3, period_seconds=60)
            self.assertTrue(allowed)
        allowed, retry_after = check_rate_limit(request, "test_action", max_actions=3, period_seconds=60)
        self.assertFalse(allowed)
        self.assertGreater(retry_after, 0)


class FullAppSmokeTest(TestCase):
    """Каждый GET-URL всех приложений должен рендериться без ошибок 500.

    Создаёт минимальные объекты нужных моделей и обходит все страницы.
    POST-эндпоинты (удаление, переключение статуса, голосование и т.д.)
    не проверяются здесь — для них есть отдельные тесты.
    """

    def setUp(self):
        self.user = User.objects.create_user("full_smoke", password="pass1234")
        self.client.force_login(self.user)

        self.project = Project.objects.create(owner=self.user, name="Smoke")
        self.task = Task.objects.create(
            owner=self.user,
            name="Active",
            project=self.project,
            deadline=timezone.localdate() + timezone.timedelta(days=7),
            priority=Task.Priority.HIGH,
        )
        self.task_done = Task.objects.create(
            owner=self.user,
            name="Done",
            project=self.project,
            status=Task.Status.DONE,
        )

        # Файлы задач и проекта (создаём реальный файл на диске)
        content = b"hello"
        self.task_file = TaskFile.objects.create(
            task=self.task,
            file=SimpleUploadedFile("test.txt", content),
            original_name="test.txt",
        )
        self.project_file = TaskFile.objects.create(
            project=self.project,
            file=SimpleUploadedFile("proj.txt", content),
            original_name="proj.txt",
        )

        self.note = Note.objects.create(owner=self.user, title="N", text="body")
        self.journal_entry = JournalEntry.objects.create(
            owner=self.user, text="Work done",
        )

        self.template = ProjectTemplate.objects.create(owner=self.user, name="Tpl")
        self.template_task = TemplateTask.objects.create(
            template=self.template, name="Tpl Task",
        )

        self.completed_project = Project.objects.create(
            owner=self.user, name="Closed",
            status=Project.Status.COMPLETED, completed_at=timezone.now(),
        )

        self.public_task = Task.objects.create(
            owner=self.user, name="Pub", share_code="pub123abc",
        )

        from meetings.models import Poll
        self.poll = Poll.objects.create(
            owner=self.user, title="Poll", organizer="Me",
            dates=["2026-08-10"], time_from=9, time_to=17,
        )

        from votes.models import Board
        self.board = Board.objects.create(
            owner=self.user, title="Board", organizer="Me",
        )

        from agenda.models import Meeting, MeetingOutcome
        self.agenda_meeting = Meeting.objects.create(
            owner=self.user, title="Meeting", organizer="Me",
        )
        self.outcome = MeetingOutcome.objects.create(
            meeting=self.agenda_meeting, title="Outcome",
        )

        from focus.models import WorkSession, TaskWorkRecord
        self.ws = WorkSession.objects.create(
            user=self.user, energy=2, focus=2, available_time=2,
        )
        self.work_record = TaskWorkRecord.objects.create(
            user=self.user, task=self.task, work_session=self.ws,
            started_at=timezone.now(),
        )

    def _get(self, url, **kwargs):
        return self.client.get(url, **kwargs)

    def _ok(self, url, **kwargs):
        """GET-запрос должен вернуть 200 (допускается редирект на логин)."""
        r = self._get(url, **kwargs)
        self.assertNotEqual(r.status_code, 500, f"GET {url} — server error")
        self.assertIn(r.status_code, (200, 302), f"GET {url} — {r.status_code}")

    # ── tasks ────────────────────────────────────────────────────────────

    def test_tasks_main_pages(self):
        today = timezone.localdate()
        urls = [
            reverse("tasks:home"),
            reverse("tasks:team_home"),
            reverse("tasks:workspace"),
            reverse("tasks:overdue_tasks"),
            reverse("tasks:completed_tasks"),
            reverse("tasks:completed_projects"),
            reverse("tasks:history"),
            reverse("tasks:stats"),
            reverse("tasks:journal"),
            reverse("tasks:journal_summary"),
            reverse("tasks:journal_summary_md"),
            reverse("tasks:calendar"),
            reverse("tasks:calendar_day",
                     kwargs={"year": today.year, "month": today.month, "day": today.day}),
            reverse("tasks:template_list"),
            reverse("tasks:template_create"),
            reverse("tasks:template_detail", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_edit", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_delete", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_run", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_task_create", kwargs={"pk": self.template.pk}),
            reverse("tasks:template_task_edit",
                     kwargs={"pk": self.template.pk, "task_pk": self.template_task.pk}),
            reverse("tasks:template_task_delete",
                     kwargs={"pk": self.template.pk, "task_pk": self.template_task.pk}),
            reverse("tasks:project_list"),
            reverse("tasks:project_create"),
            reverse("tasks:project_detail", kwargs={"pk": self.project.pk}),
            reverse("tasks:project_edit", kwargs={"pk": self.project.pk}),
            reverse("tasks:project_complete", kwargs={"pk": self.project.pk}),
            reverse("tasks:project_reopen", kwargs={"pk": self.completed_project.pk}),
            reverse("tasks:import_tasks", kwargs={"pk": self.project.pk}),
            reverse("tasks:task_create"),
            reverse("tasks:task_detail", kwargs={"pk": self.task.pk}),
            reverse("tasks:task_edit", kwargs={"pk": self.task.pk}),
            reverse("tasks:task_toggle", kwargs={"pk": self.task.pk}),
            reverse("tasks:note_list"),
            reverse("tasks:note_create"),
            reverse("tasks:note_edit", kwargs={"pk": self.note.pk}),
            reverse("tasks:journal_entry_edit", kwargs={"pk": self.journal_entry.pk}),
        ]
        for url in urls:
            with self.subTest(url=url):
                self._ok(url)

    def test_tasks_delete_confirm_pages(self):
        """DeleteView показывает страницу подтверждения (GET 200)."""
        for url in [
            reverse("tasks:task_delete", kwargs={"pk": self.task.pk}),
            reverse("tasks:project_delete", kwargs={"pk": self.project.pk}),
            reverse("tasks:template_delete", kwargs={"pk": self.template.pk}),
            reverse("tasks:note_delete", kwargs={"pk": self.note.pk}),
            reverse("tasks:journal_entry_delete", kwargs={"pk": self.journal_entry.pk}),
        ]:
            with self.subTest(url=url):
                r = self._get(url)
                self.assertEqual(r.status_code, 200, f"GET {url}")

    def test_project_complete_done_redirects(self):
        r = self._get(
            reverse("tasks:project_complete", kwargs={"pk": self.completed_project.pk})
        )
        self.assertEqual(r.status_code, 302)

    def test_project_reopen_active_redirects(self):
        r = self._get(
            reverse("tasks:project_reopen", kwargs={"pk": self.project.pk})
        )
        self.assertEqual(r.status_code, 302)

    def test_task_file_download(self):
        r = self._get(reverse("tasks:task_file_download",
                              kwargs={"pk": self.task.pk, "file_pk": self.task_file.pk}))
        self.assertIn(r.status_code, (200, 404))

    def test_project_file_download(self):
        r = self._get(reverse("tasks:project_file_download",
                              kwargs={"pk": self.project.pk, "file_pk": self.project_file.pk}))
        self.assertIn(r.status_code, (200, 404))

    def test_public_task_page(self):
        self.client.logout()
        r = self._get(reverse("tasks:public_task", kwargs={"code": self.public_task.share_code}))
        self.assertEqual(r.status_code, 200)

    def test_register_and_login_pages(self):
        self.client.logout()
        for url in [reverse("tasks:register"), reverse("tasks:login")]:
            with self.subTest(url=url):
                r = self._get(url)
                self.assertEqual(r.status_code, 200)

    # ── meetings ─────────────────────────────────────────────────────────

    def test_meetings_pages(self):
        for url in [
            reverse("meetings:home"),
            reverse("meetings:poll", kwargs={"share_code": self.poll.share_code}),
            reverse("meetings:admin", kwargs={"admin_code": self.poll.admin_code}),
        ]:
            with self.subTest(url=url):
                self._ok(url)

    # ── votes ────────────────────────────────────────────────────────────

    def test_votes_pages(self):
        for url in [
            reverse("votes:home"),
            reverse("votes:board", kwargs={"share_code": self.board.share_code}),
            reverse("votes:admin", kwargs={"admin_code": self.board.admin_code}),
            reverse("votes:protocol_md", kwargs={"admin_code": self.board.admin_code}),
        ]:
            with self.subTest(url=url):
                self._ok(url)

    # ── agenda ───────────────────────────────────────────────────────────

    def test_agenda_pages(self):
        for url in [
            reverse("agenda:home"),
            reverse("agenda:meeting", kwargs={"share_code": self.agenda_meeting.share_code}),
            reverse("agenda:admin", kwargs={"admin_code": self.agenda_meeting.admin_code}),
            reverse("agenda:outcome_detail",
                     kwargs={"admin_code": self.agenda_meeting.admin_code, "pk": self.outcome.pk}),
            reverse("agenda:outcome_edit",
                     kwargs={"admin_code": self.agenda_meeting.admin_code, "pk": self.outcome.pk}),
        ]:
            with self.subTest(url=url):
                self._ok(url)

    # ── focus ────────────────────────────────────────────────────────────

    def test_focus_pages(self):
        for url in [
            reverse("focus:dashboard"),
            reverse("focus:assess"),
            reverse("focus:history"),
            reverse("focus:statistics"),
            reverse("focus:in_progress", kwargs={"pk": self.work_record.pk}),
            reverse("focus:finish", kwargs={"pk": self.work_record.pk}),
        ]:
            with self.subTest(url=url):
                self._ok(url)

    def test_focus_recommendation_redirects_without_session(self):
        """recommendation редиректит на assess, если сессия не задана."""
        r = self._get(reverse("focus:recommendation"))
        self.assertEqual(r.status_code, 302)
